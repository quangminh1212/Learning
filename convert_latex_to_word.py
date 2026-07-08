import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def convert_latex_to_word():
    # Đường dẫn file LaTeX
    latex_file = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\main.tex'
    latex_dir = os.path.dirname(latex_file)
    
    # Tạo document Word mới
    doc = Document()
    
    # Thiết lập font Times New Roman
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    
    # Đọc file main.tex
    with open(latex_file, 'r', encoding='utf-8') as f:
        latex_content = f.read()
    
    # Thêm tiêu đề
    title = doc.add_heading('BÀI TẬP LỚN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    topic = doc.add_heading('QUẢN LÝ KHO HÀNG', level=2)
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    author = doc.add_paragraph('Nhóm thực hiện')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date = doc.add_paragraph('Hà Nội, tháng 7 năm 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Đọc và xử lý các chapter files
    chapters = [
        'chapter1_khao_sat_hien_trang',
        'chapter2_mo_ta_nghiep_vu', 
        'chapter3_phan_tich_chuc_nang',
        'chapter4_phan_tich_hanh_vi',
        'chapter5_phan_tich_tuong_tac',
        'chapter6_thiet_ke_lop_chi_tiet',
        'chapter7_thiet_ke_csdl_va_giao_dien'
    ]
    
    for chapter_file in chapters:
        chapter_path = os.path.join(latex_dir, 'chapters', f'{chapter_file}.tex')
        
        if os.path.exists(chapter_path):
            # print(f"Processing {chapter_file}...")
            process_chapter(doc, chapter_path)
        else:
            # print(f"Chapter file not found: {chapter_path}")
            pass
    
    # Lưu file Word
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\BaoCao_QuanLyKho.docx'
    doc.save(output_path)
    print("Word document created successfully")

def process_chapter(doc, chapter_path):
    """Xử lý từng chapter file"""
    with open(chapter_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Xử lý nội dung LaTeX
    lines = content.split('\n')
    
    current_section = ""
    in_itemize = False
    in_enumerate = False
    
    for line in lines:
        line = line.strip()
        
        # Bỏ qua các dòng LaTeX command
        if line.startswith('\\') and not line.startswith('\\chapter') and not line.startswith('\\section') and not line.startswith('\\subsection'):
            continue
        
        # Xử lý chapter
        if line.startswith('\\chapter'):
            chapter_title = extract_latex_title(line, '\\chapter')
            if chapter_title:
                heading = doc.add_heading(chapter_title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue
        
        # Xử lý section
        if line.startswith('\\section'):
            section_title = extract_latex_title(line, '\\section')
            if section_title:
                heading = doc.add_heading(section_title, level=2)
            continue
        
        # Xử lý subsection
        if line.startswith('\\subsection'):
            subsection_title = extract_latex_title(line, '\\subsection')
            if subsection_title:
                heading = doc.add_heading(subsection_title, level=3)
            continue
        
        # Xử lý itemize
        if '\\begin{itemize}' in line:
            in_itemize = True
            continue
        if '\\end{itemize}' in line:
            in_itemize = False
            continue
        
        # Xử lý enumerate
        if '\\begin{enumerate}' in line:
            in_enumerate = True
            continue
        if '\\end{enumerate}' in line:
            in_enumerate = False
            continue
        
        # Xử lý item
        if line.startswith('\\item'):
            item_text = line.replace('\\item', '').strip()
            if in_enumerate:
                para = doc.add_paragraph(item_text, style='List Number')
            else:
                para = doc.add_paragraph(item_text, style='List Bullet')
            continue
        
        # Xử lý text thường
        if line and not line.startswith('%') and not line.startswith('\\'):
            # Xử lý các ký tự LaTeX đặc biệt
            text = clean_latex_text(line)
            if text:
                para = doc.add_paragraph(text)

def extract_latex_title(line, command):
    """Trích xuất tiêu đề từ LaTeX command"""
    # Tìm nội dung giữa { }
    match = re.search(r'\{([^}]+)\}', line)
    if match:
        return match.group(1)
    return ""

def clean_latex_text(text):
    """Làm sạch text từ LaTeX"""
    # Thay thế các ký tự LaTeX đặc biệt
    text = text.replace('\\$', '$')
    text = text.replace('\\%', '%')
    text = text.replace('\\&', '&')
    text = text.replace('\\_', '_')
    text = text.replace('\\{', '{')
    text = text.replace('\\}', '}')
    text = text.replace('\\"', '"')
    text = text.replace("\\'", "'")
    text = text.replace('\\^', '^')
    text = text.replace('\\~', '~')
    text = text.replace('``', '"')
    text = text.replace("''", '"')
    text = text.replace('\\ldots', '...')
    text = text.replace('\\LaTeX', 'LaTeX')
    
    # Xử lý các command LaTeX khác
    text = re.sub(r'\\[a-zA-Z]+\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\[a-zA-Z]+', '', text)
    
    return text.strip()

if __name__ == '__main__':
    convert_latex_to_word()
