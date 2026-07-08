from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import zipfile
import tempfile
import shutil

def extract_images_from_docx(doc_path, output_folder):
    """Trích xuất hình ảnh từ file docx"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        extracted_images = {}
        for img_file in image_files:
            img_data = zip_ref.read(img_file)
            img_name = os.path.basename(img_file)
            img_path = os.path.join(output_folder, img_name)
            
            with open(img_path, 'wb') as f:
                f.write(img_data)
            
            extracted_images[img_name] = img_path
    
    return extracted_images

def create_complete_word_document():
    """Tạo file Word đầy đủ cho dự án du học sinh"""
    
    # Đường dẫn file
    reference_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx'
    current_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh.docx'
    
    # Thư mục tạm
    temp_folder = tempfile.mkdtemp()
    
    try:
        # Trích xuất hình ảnh từ tài liệu mẫu
        # print("Extracting images from reference document...")
        ref_images_folder = os.path.join(temp_folder, 'reference_images')
        ref_images = extract_images_from_docx(reference_doc, ref_images_folder)
        # print(f"Extracted {len(ref_images)} images")
        
        # Đọc tài liệu mẫu để lấy cấu trúc đầy đủ
        # print("Reading reference document structure...")
        ref_doc = Document(reference_doc)
        
        # Đọc tài liệu hiện tại của Nhóm 3
        # print("Reading current document...")
        curr_doc = Document(current_doc)
        
        # Tạo document mới hoàn chỉnh
        # print("Creating complete document...")
        new_doc = Document()
        
        # Thiết lập font
        new_doc.styles['Normal'].font.name = 'Times New Roman'
        new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        
        # Thêm trang bìa
        add_cover_page(new_doc)
        
        # Thêm PHÂN CÔNG CÔNG VIỆC từ tài liệu mẫu
        add_phan_cong_cong_viec(new_doc, ref_doc)
        
        # Thêm nội dung từ tài liệu hiện tại của Nhóm 3
        add_current_content(new_doc, curr_doc)
        
        # Thêm phần thiết kế hệ thống
        add_system_design_section(new_doc)
        
        # Thêm hình ảnh từ tài liệu mẫu
        add_images_from_reference(new_doc, ref_images, ref_doc)
        
        # Lưu document
        output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
        new_doc.save(output_path)
        print("Complete document created successfully")
        
        return output_path
        
    finally:
        shutil.rmtree(temp_folder)

def add_cover_page(doc):
    """Thêm trang bìa"""
    # Trang bìa
    title = doc.add_heading('PHÂN TÍCH & THIẾT KẾ', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('HỆ THỐNG QUẢN LÝ DU HỌC SINH BẬC ĐẠI HỌC CỦA TRUNG TÂM TƯ VẤN DU HỌC FLYHIGH', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    group = doc.add_paragraph('Nhóm 3')
    group.alignment = WD_ALIGN_PARAGRAPH.CENTER
    group.runs[0].bold = True
    
    doc.add_page_break()

def add_phan_cong_cong_viec(new_doc, ref_doc):
    """Thêm phần PHÂN CÔNG CÔNG VIỆC từ tài liệu mẫu"""
    # Tìm phần PHÂN CÔNG CÔNG VIỆC trong tài liệu mẫu
    start_idx = None
    end_idx = None
    
    for i, para in enumerate(ref_doc.paragraphs):
        if 'PHÂN CÔNG CÔNG VIỆC' in para.text:
            start_idx = i
        if start_idx is not None and 'Mô tả nghiệp vụ' in para.text:
            end_idx = i
            break
    
    if start_idx is not None and end_idx is not None:
        # Thêm tiêu đề
        heading = new_doc.add_heading('PHÂN CÔNG CÔNG VIỆC', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm nội dung
        for i in range(start_idx, end_idx):
            para = ref_doc.paragraphs[i]
            text = para.text.strip()
            
            if text:
                new_para = new_doc.add_paragraph(text)
                # Copy định dạng
                if para.style.name.startswith('Heading'):
                    new_para.style = para.style
                elif para.runs and para.runs[0].bold:
                    new_para.runs[0].bold = True
        
        new_doc.add_page_break()

def add_current_content(new_doc, curr_doc):
    """Thêm nội dung từ tài liệu hiện tại của Nhóm 3"""
    # Thêm PHẦN 1: TỔNG QUAN
    heading = new_doc.add_heading('PHẦN 1: TỔNG QUAN', level=1)
    
    # Copy nội dung từ tài liệu hiện tại
    for para in curr_doc.paragraphs:
        text = para.text.strip()
        if text and not text.startswith('PHÂN TÍCH & THIẾT KẾ'):
            new_para = new_doc.add_paragraph(text)
            if para.style.name.startswith('Heading'):
                new_para.style = para.style
            elif para.runs and para.runs[0].bold:
                new_para.runs[0].bold = True
    
    new_doc.add_page_break()

def add_system_design_section(new_doc):
    """Thêm phần thiết kế hệ thống"""
    # PHẦN 2: PHÂN TÍCH YÊU CẦU
    heading = new_doc.add_heading('PHẦN 2: PHÂN TÍCH YÊU CẦU', level=1)
    
    # Mô tả nghiệp vụ (đã có trong tài liệu hiện tại)
    # Phần này đã được thêm từ tài liệu hiện tại
    
    # PHẦN 3: THIẾT KẾ HỆ THỐNG
    heading = new_doc.add_heading('PHẦN 3: THIẾT KẾ HỆ THỐNG', level=1)
    
    # Biểu đồ usecase tổng quát
    uc_heading = new_doc.add_heading('Biểu đồ usecase tổng quát', level=2)
    uc_desc = new_doc.add_paragraph('Biểu đồ usecase tổng quát thể hiện các chức năng chính của hệ thống và các actor tương tác.')
    uc_note = new_doc.add_paragraph('[Hình ảnh: Biểu đồ usecase tổng quát]')
    uc_note.italic = True
    
    # Biểu đồ usecase phân rã
    uc_detail_heading = new_doc.add_heading('Biểu đồ usecase phân rã', level=2)
    uc_detail_desc = new_doc.add_paragraph('Biểu đồ usecase phân rã chi tiết hóa các usecase thành các usecase con cụ thể.')
    uc_detail_note = new_doc.add_paragraph('[Hình ảnh: Biểu đồ usecase phân rã]')
    uc_detail_note.italic = True
    
    # Đặc tả usecase
    usecase_spec_title = new_doc.add_heading('Đặc tả usecase', level=2)
    
    usecases = [
        ("UC01: Quản lý account nhân viên", "Quản lý thông tin tài khoản nhân viên hệ thống"),
        ("UC02: Đăng nhập", "Xác thực người dùng truy cập hệ thống"),
        ("UC03: Tạo/Cập nhật account sinh viên", "Quản lý tài khoản sinh viên trong hệ thống"),
        ("UC04: Tạo/Cập nhật thông tin người thân", "Quản lý thông tin liên hệ khẩn cấp của sinh viên"),
        ("UC05: Cập nhật thông tin visa", "Theo dõi và cập nhật tình trạng visa sinh viên"),
        ("UC06: Tìm kiếm sinh viên", "Tìm kiếm thông tin sinh viên theo các tiêu chí"),
        ("UC07: Tạo/Cập nhật trường học", "Quản lý thông tin các trường đại học đối tác"),
        ("UC08: Tìm kiếm trường học", "Tìm kiếm thông tin trường học theo tiêu chí"),
        ("UC09: Tạo/Cập nhật hồ sơ học tập", "Quản lý hồ sơ học thuật của sinh viên"),
        ("UC10: Cập nhật kết quả GPA hàng kỳ", "Cập nhật điểm số và kết quả học tập"),
        ("UC11: Tạo/Cập nhật học bổng", "Quản lý thông tin học bổng sinh viên"),
        ("UC12: Cập nhật chi phí tài chính từng kỳ", "Theo dõi chi phí tài chính theo kỳ học"),
        ("UC13: Xem thống kê", "Xem các báo cáo và thống kê hệ thống")
    ]
    
    for uc_name, uc_desc in usecases:
        uc_heading = new_doc.add_heading(uc_name, level=3)
        uc_heading.runs[0].bold = True
        uc_para = new_doc.add_paragraph(uc_desc)
    
    # Phân tích cấu trúc
    structure_title = new_doc.add_heading('Phân tích cấu trúc', level=2)
    structure_desc = new_doc.add_paragraph('Phân tích cấu trúc hệ thống sử dụng biểu đồ lớp (class diagram) để mô tả các thực thể và mối quan hệ.')
    
    class_diagrams = [
        "Biểu đồ lớp trong ca sử dụng UC01: Quản lý account nhân viên",
        "Biểu đồ lớp trong ca sử dụng UC02: Đăng nhập",
        "Biểu đồ lớp trong ca sử dụng UC03: Tạo/cập nhật account sinh viên",
        "Biểu đồ lớp trong ca sử dụng UC04: Tạo/cập nhập thông tin người thân",
        "Biểu đồ lớp trong ca sử dụng UC05: Cập nhật thông tin visa",
        "Biểu đồ lớp trong ca sử dụng UC06: Tìm kiếm sinh viên",
        "Biểu đồ lớp trong ca sử dụng UC07: Tạo/Cập nhật trường học",
        "Biểu đồ lớp trong ca sử dụng UC08: Tìm kiếm trường học",
        "Biểu đồ lớp trong ca sử dụng UC09: Tạo/Cập nhật hồ sơ học tập",
        "Biểu đồ lớp trong ca sử dụng UC10: Cập nhật kết quả GPA hàng kỳ",
        "Biểu đồ lớp trong ca sử dụng UC11: Tạo/Cập nhật học bổng",
        "Biểu đồ lớp trong ca sử dụng UC12: Cập nhật chi phí tài chính hàng kỳ",
        "Biểu đồ lớp trong ca sử dụng UC13: Xem thống kê"
    ]
    
    for diagram_name in class_diagrams:
        diag_heading = new_doc.add_heading(diagram_name, level=3)
        diag_heading.runs[0].bold = True
        diag_note = new_doc.add_paragraph(f'[Hình ảnh: {diagram_name}]')
        diag_note.italic = True
    
    # Phân tích hành vi
    behavior_title = new_doc.add_heading('Phân tích hành vi', level=2)
    behavior_desc = new_doc.add_paragraph('Phân tích hành vi sử dụng biểu đồ trình tự (sequence diagram) để mô tả luồng tác động giữa các đối tượng.')
    
    sequence_diagrams = [
        "Biểu đồ trình tự cho ca sử dụng UC01: Quản lý account nhân viên",
        "Biểu đồ trình tự cho ca sử dụng UC02: Đăng nhập",
        "Biểu đồ trình tự cho ca sử dụng UC03: Tạo/Chỉnh sửa account sinh viên",
        "Biểu đồ trình tự cho ca sử dụng UC04: Tạo/Cập nhật thông tin người thân",
        "Biểu đồ trình tự cho ca sử dụng UC05: Cập nhật thông tin visa",
        "Biểu đồ trình tự cho ca sử dụng UC06: Tìm kiếm sinh viên",
        "Biểu đồ trình tự cho ca sử dụng UC07: Tạo/Cập nhật trường học",
        "Biểu đồ trình tự cho ca sử dụng UC08: Tìm kiếm trường học",
        "Biểu đồ trình tự cho ca sử dụng UC09: Tạo/Cập nhật hồ sơ học tập",
        "Biểu đồ trình tự cho ca sử dụng UC10: Cập nhật kết quả GPA hàng kỳ",
        "Biểu đồ trình tự cho ca sử dụng UC11: Tạo/Cập nhật học bổng",
        "Biểu đồ trình tự cho ca sử dụng UC12: Cập nhật chi phí tài chính từng kỳ",
        "Biểu đồ trình tự cho ca sử dụng UC13: Xem thống kê"
    ]
    
    for seq_name in sequence_diagrams:
        seq_heading = new_doc.add_heading(seq_name, level=3)
        seq_heading.runs[0].bold = True
        seq_note = new_doc.add_paragraph(f'[Hình ảnh: {seq_name}]')
        seq_note.italic = True

def add_images_from_reference(new_doc, ref_images, ref_doc):
    """Thêm hình ảnh từ tài liệu mẫu"""
    image_list = list(ref_images.values())
    
    # Thêm hình ảnh vào cuối document
    for i, img_path in enumerate(image_list):
        try:
            para = new_doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(img_path, width=Inches(6.0))
        except Exception as e:
            pass

if __name__ == '__main__':
    create_complete_word_document()
