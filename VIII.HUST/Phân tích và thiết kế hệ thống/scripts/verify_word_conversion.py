from docx import Document

def verify_word_document():
    doc_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\BaoCao_QuanLyKho.docx'
    
    doc = Document(doc_path)
    
    # Đếm paragraphs
    paragraph_count = len(doc.paragraphs)
    
    # Đếm headings
    heading_count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            heading_count += 1
    
    # Lấy nội dung mẫu
    content_sample = []
    for i, para in enumerate(doc.paragraphs[:20]):
        if para.text.strip():
            content_sample.append(f"{i+1}: {para.text[:80]}")
    
    # Kiểm tra các chapter quan trọng
    full_text = '\n'.join([para.text for para in doc.paragraphs])
    
    important_keywords = {
        'CHƯƠNG 1': 'CHƯƠNG 1' in full_text or 'Chương 1' in full_text,
        'KHẢO SÁT': 'KHẢO SÁT' in full_text or 'Khảo sát' in full_text,
        'MÔ TẢ NGHIỆP VỤ': 'MÔ TẢ NGHIỆP VỤ' in full_text or 'Mô tả nghiệp vụ' in full_text,
        'PHÂN TÍCH CHỨC NĂNG': 'PHÂN TÍCH CHỨC NĂNG' in full_text or 'Phân tích chức năng' in full_text,
        'THIẾT KẾ': 'THIẾT KẾ' in full_text or 'Thiết kế' in full_text,
    }
    
    with open(r'C:\Dev\Learning\word_conversion_verification.txt', 'w', encoding='utf-8') as f:
        f.write('WORD CONVERSION VERIFICATION\n')
        f.write('============================\n\n')
        f.write(f'Total paragraphs: {paragraph_count}\n')
        f.write(f'Total headings: {heading_count}\n\n')
        
        f.write('IMPORTANT KEYWORDS:\n')
        for keyword, present in important_keywords.items():
            status = 'PRESENT' if present else 'MISSING'
            f.write(f'{keyword}: {status}\n')
        
        f.write('\nCONTENT SAMPLE (First 20 lines):\n')
        for line in content_sample:
            f.write(line + '\n')
    
    print('Verification completed!')
    print('Results saved to word_conversion_verification.txt')
    print(f'Total paragraphs: {paragraph_count}')
    print(f'Total headings: {heading_count}')

if __name__ == '__main__':
    verify_word_document()
