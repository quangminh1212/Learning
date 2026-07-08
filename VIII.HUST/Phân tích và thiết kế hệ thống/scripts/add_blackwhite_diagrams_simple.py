from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_blackwhite_diagrams():
    """Thêm placeholder cho diagram đen trắng vào file Word"""
    
    word_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
    tex_folder = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\diagrams\tex'
    
    doc = Document(word_doc)
    
    # Lấy danh sách file .tex
    tex_files = [f for f in os.listdir(tex_folder) if f.endswith('.tex')]
    tex_files.sort()
    
    # print(f"Found {len(tex_files)} diagram files")
    
    # Thêm section diagram
    diagram_heading = doc.add_heading('CÁC BIỂU ĐỒ THIẾT KẾ HỆ THỐNG (MÀU ĐEN)', level=1)
    diagram_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    note = doc.add_paragraph('Lưu ý: Các diagram dưới đây cần được tạo theo phong cách đen trắng (grayscale)')
    note.italic = True
    
    # Phân loại diagram
    categories = {
        'ACTIVITY DIAGRAMS': [f for f in tex_files if 'activity' in f],
        'BUBBLE FLOW DIAGRAM': [f for f in tex_files if 'bfd' in f],
        'DATA FLOW DIAGRAMS': [f for f in tex_files if 'dfd' in f],
        'USECASE DIAGRAMS': [f for f in tex_files if 'usecase' in f],
        'CLASS DIAGRAMS': [f for f in tex_files if 'class' in f],
        'SEQUENCE DIAGRAMS': [f for f in tex_files if 'sequence' in f],
        'STATE DIAGRAMS': [f for f in tex_files if 'state' in f],
        'COMPONENT DIAGRAM': [f for f in tex_files if 'component' in f],
        'ENTITY RELATIONSHIP DIAGRAMS': [f for f in tex_files if 'erd' in f],
    }
    
    for category, files in categories.items():
        if files:
            heading = doc.add_heading(category, level=2)
            for tex_file in files:
                diagram_name = tex_file.replace('.tex', '').replace('_', ' ').upper()
                diagram_heading = doc.add_heading(diagram_name, level=3)
                
                # Thêm hướng dẫn tạo diagram đen trắng
                instructions = doc.add_paragraph()
                instructions.add_run('File nguồn: ').bold = True
                instructions.add_run(tex_file)
                
                bw_note = doc.add_paragraph('Tạo diagram đen trắng từ file LaTeX:')
                bw_note.italic = True
                
                # Thêm các bước
                steps = [
                    '1. Mở file .tex trong thư mục diagrams/tex',
                    '2. Thêm dòng sau vào preamble: \\pgfsetfillcolor{black}',
                    '3. Thêm dòng sau vào preamble: \\pgfsetstrokecolor{black}',
                    '4. Biên dịch với pdflatex để tạo PDF đen trắng',
                    '5. Chuyển PDF thành hình ảnh và thêm vào đây'
                ]
                
                for step in steps:
                    step_para = doc.add_paragraph(step, style='List Bullet')
                    step_para.runs[0].font.size = Pt(10)
                
                doc.add_paragraph()
    
    # Thêm các diagram mẫu từ tài liệu mẫu (đen trắng)
    add_sample_diagrams_from_reference(doc)
    
    # Lưu file
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_BLACKWHITE.docx'
    doc.save(output_path)
    print("Document created successfully")

def add_sample_diagrams_from_reference(doc):
    """Thêm các diagram mẫu từ tài liệu mẫu (đen trắng)"""
    
    reference_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx'
    
    try:
        ref_doc = Document(reference_doc)
        
        # Thêm section diagram mẫu
        sample_heading = doc.add_heading('CÁC BIỂU ĐỒ MẪU TỪ TÀI LIỆU CHUẨN', level=1)
        sample_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Trích xuất hình ảnh từ tài liệu mẫu
        import zipfile
        import tempfile
        
        temp_folder = tempfile.mkdtemp()
        
        try:
            # Trích xuất hình ảnh
            with zipfile.ZipFile(reference_doc, 'r') as zip_ref:
                image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
                
                extracted_count = 0
                for img_file in image_files[:20]:  # Giới hạn 20 hình ảnh đầu tiên
                    img_data = zip_ref.read(img_file)
                    img_name = os.path.basename(img_file)
                    img_path = os.path.join(temp_folder, img_name)
                    
                    with open(img_path, 'wb') as f:
                        f.write(img_data)
                    
                    # Thêm vào Word
                    try:
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(img_path, width=Inches(5.0))
                        extracted_count += 1
                    except:
                        pass
                
            # print(f"Added {extracted_count} sample diagrams from reference")
            
        finally:
            import shutil
            shutil.rmtree(temp_folder)
            
    except Exception as e:
        # print(f"Could not add sample diagrams: {e}")
        pass

if __name__ == '__main__':
    add_blackwhite_diagrams()
