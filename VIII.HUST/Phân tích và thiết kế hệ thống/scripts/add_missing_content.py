from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_missing_content():
    # Đọc báo cáo hiện tại
    current_doc = Document(r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh.docx')
    
    # Thêm PHÂN CÔNG CÔNG VIỆC vào đầu
    title = current_doc.add_heading('PHÂN CÔNG CÔNG VIỆC', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Nội dung phân công công việc
    pc_content = [
        "",
        "Khảo sát hiện trạng",
        "",
        "Giới thiệu doanh nghiệp",
        "Trung tâm tư vấn du học FlyHigh là một đơn vị hoạt động trong lĩnh vực tư vấn và hỗ trợ du học sinh, chuyên cung cấp các dịch vụ: tư vấn lộ trình, xử lý hồ sơ nhập học, xin visa, săn học bổng và hỗ trợ sinh viên trong suốt quá trình học tập tại nước ngoài (Hàn Quốc, Nhật Bản, Mỹ, Úc, châu Âu...).",
        "",
        "Với số lượng học viên ngày càng tăng, khối lượng thông tin cần quản lý (hồ sơ cá nhân, tiến độ visa, tài chính, thông tin trường đối tác) trở nên cực kỳ lớn và phức tạp. Việc chuyển đổi từ quản lý truyền thống sang một hệ thống phần mềm chuyên biệt là nhu cầu cấp thiết để tối ưu hóa vận hành.",
        "",
        "Thực trạng quy trình quản lý hiện tại và các hạn chế",
        "Hiện nay, trung tâm FlyHigh vẫn đang vận hành và quản lý dữ liệu du học sinh chủ yếu thông qua các công cụ văn phòng rời rạc như Google Sheets, Excel kết hợp với việc lưu trữ hồ sơ giấy tại văn phòng. Thực trạng này đang dẫn đến nhiều bất cập cụ thể như sau:",
        "",
        "Quản lý thông tin học viên phân tán: Mỗi tư vấn viên (Sale/Tư vấn) tự quản lý một file Excel riêng về danh sách khách hàng của mình. Khi học viên chuyển sang giai đoạn xử lý hồ sơ (Bộ phận Hồ sơ) hoặc theo dõi visa, dữ liệu lại được copy thủ công sang một file khác. Điều này dẫn đến tình trạng tam sao thất bản, sai lệch thông tin (sai số điện thoại, viết sai tên tiếng Anh, nhầm ngày sinh).",
        "",
        "Khó khăn trong việc theo dõi hạn chót: Các thông tin có tính chất thời hạn như hạn nộp hồ sơ trường (Application Deadline), hạn đóng học phí, ngày hết hạn visa, hạn bảo hiểm hiện đang được theo dõi thủ công bằng lịch cá nhân hoặc trí nhớ của nhân viên. Trung tâm đã từng gặp sự cố chậm trễ tiến độ nộp hồ sơ của học viên do không có hệ thống cảnh báo tự động.",
        "",
        "Theo dõi tài chính nhập nhằng: Việc quản lý các khoản phí (phí dịch vụ trung tâm, học phí đặt cọc cho trường, phí bảo hiểm, phí visa) được ghi nhận thủ công trên Excel và đối chiếu với kế toán qua biên lai giấy. Khi học viên hoặc phụ huynh hỏi về công nợ hoặc tiến độ thanh toán, nhân viên mất nhiều thời gian để tra cứu và phản hồi.",
        "",
        "Mất kết nối thông tin sau khi bay: Sau khi học viên xuất cảnh, thông tin về tình trạng học tập, gia hạn visa tại nước ngoài hoặc tình hình nhận học bổng các kỳ tiếp theo hầu như không được cập nhật tập trung. Trung tâm chỉ liên lạc qua Zalo/Facebook khi có sự cố phát sinh, gây khó khăn cho việc duy trì mạng lưới cựu du học sinh.",
        "",
        "Công tác báo cáo, thống kê chậm trễ: Ban giám đốc không có cái nhìn tổng quan theo thời gian thực (Real-time) về: Tỷ lệ đỗ visa tháng này là bao nhiêu? Doanh thu dự kiến từ kỳ nhập học tới là bao nhiêu? Quốc gia nào đang có lượng học sinh đăng ký nhiều nhất? Mọi báo cáo đều phải chờ nhân viên tổng hợp thủ công từ các file Excel, mất từ 1 - 2 ngày và dễ sai sót.",
        "",
        "Mục tiêu của hệ thống",
        "Hệ thống quản lý du học sinh của trung tâm FlyHigh được xây dựng nhằm hỗ trợ việc quản lý thông tin sinh viên một cách khoa học và hiệu quả. Hệ thống giúp lưu trữ, quản lý và xử lý các dữ liệu liên quan đến du học sinh trong suốt quá trình học tập tại nước ngoài.",
        "",
        "Mục tiêu chính của hệ thống bao gồm:",
        "Tập trung hóa dữ liệu: Xóa bỏ hoàn toàn các file Excel rời rạc. Toàn bộ thông tin từ lúc học viên tư vấn, làm hồ sơ, xin visa, đóng phí cho đến khi học tập tại nước ngoài đều được lưu trữ trên một cơ sở dữ liệu duy nhất, phân quyền rõ ràng cho từng bộ phận.",
        "",
        "Số hóa hồ sơ tài chính & Học thuật: Quản lý chi tiết từng khoản thu/chi của học viên, theo dõi học bổng theo từng kỳ học, và lưu trữ lịch sử học tập (Trường, ngành, bảng điểm nếu có) một cách khoa học.",
        "",
        "Báo cáo thông minh (Smart Reporting): Cung cấp hệ thống Dashboard trực quan cho Ban quản lý, xuất báo cáo thống kê tự động về tỷ lệ đỗ visa, doanh thu, số lượng học sinh theo từng quốc gia chỉ bằng một cú click chuột.",
        "",
        "Nâng cao năng suất & Bảo mật: Giảm 80% thời gian nhập liệu trùng lặp giữa các phòng ban, bảo mật thông tin cá nhân của học viên và người thân theo đúng quy định.",
        "",
        "Thông qua hệ thống này, trung tâm FlyHigh có thể quản lý thông tin du học sinh một cách tập trung, giảm thiểu sai sót trong quá trình quản lý dữ liệu và nâng cao hiệu quả làm việc của nhân viên.",
    ]
    
    for line in pc_content:
        if line.strip():
            para = current_doc.add_paragraph(line.strip())
            if line.isupper() and len(line) < 50 and ':' not in line:
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(14)
            elif ':' in line and not line.startswith('Với') and not line.startswith('Hiện nay'):
                para.runs[0].bold = True
            elif line.startswith('-'):
                para.runs[0].bold = True
    
    # Thêm phần thiết kế vào cuối
    current_doc.add_page_break()
    
    # PHẦN 3: THIẾT KẾ HỆ THỐNG
    design_title = current_doc.add_heading('PHẦN 3: THIẾT KẾ HỆ THỐNG', level=1)
    design_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Biểu đồ usecase tổng quát
    uc_title = current_doc.add_heading('Biểu đồ usecase tổng quát', level=2)
    uc_desc = current_doc.add_paragraph("Biểu đồ usecase tổng quát thể hiện các chức năng chính của hệ thống và các actor tương tác.")
    uc_note = current_doc.add_paragraph("[Chèn hình ảnh: Biểu đồ usecase tổng quát]")
    uc_note.italic = True
    
    # Biểu đồ usecase phân rã
    uc_detail_title = current_doc.add_heading('Biểu đồ usecase phân rã', level=2)
    uc_detail_desc = current_doc.add_paragraph("Biểu đồ usecase phân rã chi tiết hóa các usecase thành các usecase con cụ thể.")
    uc_detail_note = current_doc.add_paragraph("[Chèn hình ảnh: Biểu đồ usecase phân rã]")
    uc_detail_note.italic = True
    
    # Đặc tả usecase
    usecase_spec_title = current_doc.add_heading('Đặc tả usecase', level=2)
    
    usecases = [
        ("UC01: Quản lý account nhân viên", "Quản lý thông tin tài khoản nhân viên hệ thống"),
        ("UC02: Đăng nhập", "Xác thực người dùng truy cập hệ thống"),
        ("UC03: Tạo/Cập nhật account sinh viên", "Quản lý tài khoản sinh viên trong hệ thống"),
        ("UC04: Tạo/Cập nhật thông tin người thân", "Quản lý thông tin liên hệ khẩn cấp của sinh viên"),
        ("UC05: Cập nhật thông tin visa", "Theo dõi và cập nhật tình trạng visa sinh viên"),
        ("UC06: Tìm kiếm sinh viên", "Tìm kiếm thông tin sinh viên theo các tiêu chí"),
        ("UC07: Tạo/Cập nhật trường học", "Quản lý thông tin các trường đại học đối tác"),
        ("UC08: Tìm kiếm trường học", "Tìm kiếm thông tin trường học theo tiêu chí"),
        ("UC09: Tạo/Cập nhật hồ sơ học tập", "Quản lý hồ sơ học thuật của sinh viên"),
        ("UC10: Cập nhật kết quả GPA hàng kỳ", "Cập nhật điểm số và kết quả học tập"),
        ("UC11: Tạo/Cập nhật học bổng", "Quản lý thông tin học bổng sinh viên"),
        ("UC12: Cập nhật chi phí tài chính từng kỳ", "Theo dõi chi phí tài chính theo kỳ học"),
        ("UC13: Xem thống kê", "Xem các báo cáo và thống kê hệ thống")
    ]
    
    for uc_name, uc_desc in usecases:
        uc_heading = current_doc.add_heading(uc_name, level=3)
        uc_heading.runs[0].bold = True
        uc_para = current_doc.add_paragraph(uc_desc)
    
    # Phân tích cấu trúc
    structure_title = current_doc.add_heading('Phân tích cấu trúc', level=2)
    structure_desc = current_doc.add_paragraph("Phân tích cấu trúc hệ thống sử dụng biểu đồ lớp (class diagram) để mô tả các thực thể và mối quan hệ.")
    
    class_diagrams = [
        "Biểu đồ lớp trong ca sử dụng UC01: Quản lý account nhân viên",
        "Biểu đồ lớp trong ca sử dụng UC02: Đăng nhập",
        "Biểu đồ lớp trong ca sử dụng UC03: Tạo/cập nhật account sinh viên",
        "Biểu đồ lớp trong ca sử dụng UC04: Tạo/cập nhập thông tin người thân",
        "Biểu đồ lớp trong ca sử dụng UC05: Cập nhật thông tin visa",
        "Biểu đồ lớp trong ca sử dụng UC06: Tìm kiếm sinh viên",
        "Biểu đồ lớp trong ca sử dụng UC07: Tạo/Cập nhật trường học",
        "Biểu đồ lớp trong ca sử dụng UC08: Tìm kiếm trường học",
        "Biểu đồ lớp trong ca sử dụng UC09: Tạo/Cập nhật hồ sơ học tập",
        "Biểu đồ lớp trong ca sử dụng UC10: Cập nhật kết quả GPA hàng kỳ",
        "Biểu đồ lớp trong ca sử dụng UC11: Tạo/Cập nhật học bổng",
        "Biểu đồ lớp trong ca sử dụng UC12: Cập nhật chi phí tài chính hàng kỳ",
        "Biểu đồ lớp trong ca sử dụng UC13: Xem thống kê"
    ]
    
    for diagram_name in class_diagrams:
        diag_heading = current_doc.add_heading(diagram_name, level=3)
        diag_heading.runs[0].bold = True
        diag_note = current_doc.add_paragraph(f"[Chèn hình ảnh: {diagram_name}]")
        diag_note.italic = True
    
    # Phân tích hành vi
    behavior_title = current_doc.add_heading('Phân tích hành vi', level=2)
    behavior_desc = current_doc.add_paragraph("Phân tích hành vi sử dụng biểu đồ trình tự (sequence diagram) để mô tả luồng tác động giữa các đối tượng.")
    
    sequence_diagrams = [
        "Biểu đồ trình tự cho ca sử dụng UC01: Quản lý account nhân viên",
        "Biểu đồ trình tự cho ca sử dụng UC02: Đăng nhập",
        "Biểu đồ trình tự cho ca sử dụng UC03: Tạo/Chỉnh sửa account sinh viên",
        "Biểu đồ trình tự cho ca sử dụng UC04: Tạo/Cập nhật thông tin người thân",
        "Biểu đồ trình tự cho ca sử dụng UC05: Cập nhật thông tin visa",
        "Biểu đồ trình tự cho ca sử dụng UC06: Tìm kiếm sinh viên",
        "Biểu đồ trình tự cho ca sử dụng UC07: Tạo/Cập nhật trường học",
        "Biểu đồ trình tự cho ca sử dụng UC08: Tìm kiếm trường học",
        "Biểu đồ trình tự cho ca sử dụng UC09: Tạo/Cập nhật hồ sơ học tập",
        "Biểu đồ trình tự cho ca sử dụng UC10: Cập nhật kết quả GPA hàng kỳ",
        "Biểu đồ trình tự cho ca sử dụng UC11: Tạo/Cập nhật học bổng",
        "Biểu đồ trình tự cho ca sử dụng UC12: Cập nhật chi phí tài chính từng kỳ",
        "Biểu đồ trình tự cho ca sử dụng UC13: Xem thống kê"
    ]
    
    for seq_name in sequence_diagrams:
        seq_heading = current_doc.add_heading(seq_name, level=3)
        seq_heading.runs[0].bold = True
        seq_note = current_doc.add_paragraph(f"[Chèn hình ảnh: {seq_name}]")
        seq_note.italic = True
    
    # Lưu document
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_UPDATED.docx'
    current_doc.save(output_path)
    print(f'Document updated successfully!')
    print(f'Saved as: {output_path}')
    print('Added sections:')
    print('1. PHÂN CÔNG CÔNG VIỆC (including Khảo sát hiện trạng)')
    print('2. PHẦN 3: THIẾT KẾ HỆ THỐNG')
    print('3. Biểu đồ usecase tổng quát và phân rã')
    print('4. Đặc tả usecase (UC01-UC13)')
    print('5. Phân tích cấu trúc (Biểu đồ lớp)')
    print('6. Phân tích hành vi (Biểu đồ trình tự)')

if __name__ == '__main__':
    add_missing_content()
