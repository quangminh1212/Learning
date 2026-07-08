from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_missing_sections():
    # Đọc tài liệu mẫu để lấy nội dung
    reference_doc = Document(r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx')
    
    # Đọc báo cáo hiện tại
    current_doc = Document(r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh.docx')
    
    # Trích xuất nội dung từ tài liệu mẫu
    ref_text = []
    for para in reference_doc.paragraphs:
        ref_text.append(para.text)
    
    ref_full_text = '\n'.join(ref_text)
    
    # Tìm các phần cần thêm từ tài liệu mẫu
    sections_to_add = []
    
    # PHÂN CÔNG CÔNG VIỆC
    if 'PHÂN CÔNG CÔNG VIỆC' in ref_full_text:
        start_idx = ref_full_text.index('PHÂN CÔNG CÔNG VIỆC')
        # Tìm phần tiếp theo để xác định giới hạn
        next_section_idx = ref_full_text.find('Khảo sát hiện trạng', start_idx)
        if next_section_idx != -1:
            sections_to_add.append(('PHÂN CÔNG CÔNG VIỆC', ref_full_text[start_idx:next_section_idx].strip()))
    
    # Khảo sát hiện trạng
    if 'Khảo sát hiện trạng' in ref_full_text:
        start_idx = ref_full_text.index('Khảo sát hiện trạng')
        next_section_idx = ref_full_text.find('Mô tả nghiệp vụ', start_idx)
        if next_section_idx != -1:
            sections_to_add.append(('Khảo sát hiện trạng', ref_full_text[start_idx:next_section_idx].strip()))
    
    # Thêm vào đầu document
    for section_name, section_content in reversed(sections_to_add):
        # Thêm tiêu đề section
        title = current_doc.add_heading(section_name, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm nội dung
        lines = section_content.split('\n')
        for line in lines:
            if line.strip():
                para = current_doc.add_paragraph(line.strip())
                if line.isupper() and len(line) < 50:
                    para.runs[0].bold = True
    
    # Thêm các phần thiết kế vào cuối document
    design_sections = [
        'Biểu đồ usecase tổng quát',
        'Biểu đồ usecase phân rã', 
        'Đặc tả usecase',
        'Phân tích cấu trúc',
        'Phân tích hành vi'
    ]
    
    for section_name in design_sections:
        if section_name in ref_full_text:
            start_idx = ref_full_text.find(section_name)
            # Tìm section tiếp theo
            next_section = None
            for next_sec in design_sections:
                if next_sec != section_name and next_sec in ref_full_text[start_idx + len(section_name):]:
                    next_section = next_sec
                    break
            
            if next_section:
                end_idx = ref_full_text.find(next_section, start_idx)
                section_content = ref_full_text[start_idx:end_idx].strip()
            else:
                section_content = ref_full_text[start_idx:].strip()
            
            # Thêm tiêu đề
            title = current_doc.add_heading(section_name, level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Thêm nội dung
            lines = section_content.split('\n')
            for line in lines:
                if line.strip():
                    para = current_doc.add_paragraph(line.strip())
                    # Format các dòng quan trọng
                    if any(keyword in line for keyword in ['UC01', 'UC02', 'UC03', 'UC04', 'UC05', 'UC06', 'UC07', 'UC08', 'UC09', 'UC10', 'UC11', 'UC12', 'UC13']):
                        para.runs[0].bold = True
                    if line.startswith('Biểu đồ'):
                        para.runs[0].bold = True
    
    # Lưu document
    current_doc.save(r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_UPDATED.docx')
    print('Document updated successfully!')
    print('Saved as: Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_UPDATED.docx')

if __name__ == '__main__':
    add_missing_sections()
