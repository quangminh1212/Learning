from docx import Document
from docx.shared import Inches
import os
import zipfile
import tempfile

def extract_images_from_docx(doc_path, output_folder):
    """Trích xuất hình ảnh từ file docx"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # File docx thực chất là một zip file
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        # Tìm tất cả các file hình ảnh trong thư mục media
        image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        extracted_images = []
        for img_file in image_files:
            # Trích xuất hình ảnh
            img_data = zip_ref.read(img_file)
            img_name = os.path.basename(img_file)
            img_path = os.path.join(output_folder, img_name)
            
            with open(img_path, 'wb') as f:
                f.write(img_data)
            
            extracted_images.append(img_path)
    
    return extracted_images

def get_document_structure_with_images(doc_path):
    """Lấy cấu trúc document và vị trí hình ảnh"""
    doc = Document(doc_path)
    
    structure = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        has_image = bool(para._element.xpath('.//pic:pic'))
        
        if text or has_image:
            structure.append({
                'index': i,
                'type': 'paragraph',
                'text': text,
                'has_image': has_image
            })
    
    return structure

def copy_images_from_reference_to_current():
    """Sao chép hình ảnh từ tài liệu mẫu sang tài liệu hiện tại"""
    
    # Đường dẫn file
    reference_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx'
    current_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh.docx'
    updated_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_FINAL.docx'
    
    # Thư mục tạm để trích xuất hình ảnh
    temp_folder = tempfile.mkdtemp()
    ref_images_folder = os.path.join(temp_folder, 'reference_images')
    
    print("Extracting images from reference document...")
    ref_images = extract_images_from_docx(reference_doc, ref_images_folder)
    print(f"Extracted {len(ref_images)} images from reference document")
    
    # Đọc cấu trúc tài liệu mẫu
    print("Analyzing reference document structure...")
    ref_structure = get_document_structure_with_images(reference_doc)
    
    # Đọc cấu trúc tài liệu hiện tại
    print("Analyzing current document structure...")
    curr_structure = get_document_structure_with_images(current_doc)
    
    # Mở tài liệu hiện tại để chỉnh sửa
    doc = Document(current_doc)
    
    # Tìm các vị trí cần thêm hình ảnh dựa trên cấu trúc
    # Đây là cách tiếp cận đơn giản - thêm hình ảnh vào cuối các section tương ứng
    
    # Đầu tiên, đọc tài liệu mẫu để biết hình ảnh nên đặt ở đâu
    ref_doc = Document(reference_doc)
    
    # Tạo mapping giữa text và hình ảnh trong tài liệu mẫu
    ref_image_mapping = []
    for i, para in enumerate(ref_doc.paragraphs):
        text = para.text.strip()
        has_image = bool(para._element.xpath('.//pic:pic'))
        if has_image:
            # Lấy hình ảnh từ paragraph
            for run in para.runs:
                if run._element.xpath('.//pic:pic'):
                    # Tìm hình ảnh tương ứng
                    ref_image_mapping.append({
                        'index': i,
                        'text_before': text,
                        'image_index': len(ref_image_mapping)
                    })
                    break
    
    print(f"Found {len(ref_image_mapping)} image positions in reference document")
    
    # Thêm hình ảnh vào tài liệu hiện tại
    # Cách tiếp cận đơn giản: thêm hình ảnh vào cuối document
    # Người dùng có thể di chuyển chúng đến vị trí đúng sau
    
    print("Adding images to document...")
    
    # Đảm bảo có đủ hình ảnh được thêm
    for i, img_path in enumerate(ref_images[:20]):  # Giới hạn 20 hình ảnh đầu tiên
        try:
            # Thêm một paragraph mới
            para = doc.add_paragraph()
            
            # Thêm hình ảnh
            run = para.add_run()
            try:
                run.add_picture(img_path, width=Inches(6.0))
                print(f"Added image {i+1}: {os.path.basename(img_path)}")
            except Exception as e:
                print(f"Could not add image {i+1}: {e}")
        except Exception as e:
            print(f"Error processing image {i+1}: {e}")
    
    # Lưu document
    doc.save(updated_doc)
    print(f"Document saved to: {updated_doc}")
    
    # Dọn dẹp thư mục tạm
    import shutil
    shutil.rmtree(temp_folder)
    
    return updated_doc

if __name__ == '__main__':
    result_path = copy_images_from_reference_to_current()
    print(f"\nProcess completed!")
    print(f"Updated document: {result_path}")
