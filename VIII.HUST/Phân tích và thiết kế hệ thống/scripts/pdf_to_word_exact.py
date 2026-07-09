# -*- coding: utf-8 -*-
"""
Tạo Word trình bày Y HỆT PDF: mỗi trang PDF = 1 trang Word (ảnh full page A4).
Giữ nguyên mục lục, hình vẽ, căn lề, font như main.pdf.
"""
from __future__ import annotations

import io
import tempfile
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, Twips, Emu

BASE = Path(r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao")
PDF_PATH = BASE / "main.pdf"
# File chính khuyến dùng — khớp visual PDF
OUT_EXACT = BASE / "BaoCao_QuanLyKhoHang.docx"
# Bản phụ (cùng nội dung)
OUT_MIRROR = BASE / "BaoCao_QuanLyKho_PDFMirror.docx"

# A4 in EMU (914400 EMU = 1 inch)
A4_W_CM = 21.0
A4_H_CM = 29.7
# Render scale: 2.0 ≈ 144 dpi, 2.5 ≈ 180 dpi — cân bằng nét/kích thước file
RENDER_SCALE = 2.2


def set_run_font(run, name="Times New Roman", size=11):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def suppress_header_footer_distance(section):
    """Minimize header/footer so image can fill page."""
    section.header_distance = Cm(0)
    section.footer_distance = Cm(0)


def clear_paragraph_spacing(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Remove indent
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)


def set_narrow_margins(section):
    """Zero margins — page image already includes PDF margins."""
    section.page_width = Cm(A4_W_CM)
    section.page_height = Cm(A4_H_CM)
    section.orientation = WD_ORIENT.PORTRAIT
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    suppress_header_footer_distance(section)


def remove_header_footer_content(section):
    for part in (section.header, section.footer):
        for p in part.paragraphs:
            p.clear()


def build_exact_word(pdf_path: Path, out_path: Path, scale: float = RENDER_SCALE) -> Path:
    if not pdf_path.exists():
        raise FileNotFoundError(pdf_path)

    pdf = fitz.open(pdf_path)
    n = pdf.page_count
    print(f"PDF: {pdf_path.name} — {n} pages")

    doc = Document()
    section = doc.sections[0]
    set_narrow_margins(section)
    remove_header_footer_content(section)

    # Content width/height for full-bleed image
    img_width = Cm(A4_W_CM)
    # Keep aspect: A4 ratio

    mat = fitz.Matrix(scale, scale)

    for i in range(n):
        page = pdf[i]
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")

        # Một paragraph / trang; page_break_before tránh trang trống thừa
        if i == 0 and doc.paragraphs:
            p = doc.paragraphs[0]
            p.clear()
        else:
            p = doc.add_paragraph()

        clear_paragraph_spacing(p)
        if i > 0:
            p.paragraph_format.page_break_before = True

        run = p.add_run()
        # Khóa đúng khổ A4 để không tràn trang / lệch lề
        run.add_picture(
            io.BytesIO(img_bytes),
            width=Cm(A4_W_CM),
            height=Cm(A4_H_CM),
        )

        if (i + 1) % 10 == 0 or i + 1 == n:
            print(f"  page {i + 1}/{n}")

    pdf.close()

    # Ensure all sections have zero margins
    for sec in doc.sections:
        set_narrow_margins(sec)
        remove_header_footer_content(sec)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    size_mb = out_path.stat().st_size / (1024 * 1024)
    print(f"Saved: {out_path}")
    print(f"Size: {size_mb:.2f} MB, pages: {n}")
    return out_path


def main():
    print("=== PDF → Word (exact visual mirror) ===")
    build_exact_word(PDF_PATH, OUT_EXACT)
    # Mirror copy with explicit name
    import shutil

    shutil.copy2(OUT_EXACT, OUT_MIRROR)
    print(f"Copy: {OUT_MIRROR}")
    print("Done. Mở file Word sẽ thấy từng trang y hệt main.pdf (mục lục + đủ hình).")


if __name__ == "__main__":
    main()
