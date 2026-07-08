import os
import subprocess
import tempfile
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_blackwhite_diagrams():
    """Tạo diagram đen trắng từ file LaTeX và thêm vào Word"""
    
    # Thư mục diagram LaTeX
    tex_folder = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\diagrams\tex'
    
    # Thư mục diagram PDF đen trắng
    output_folder = tempfile.mkdtemp()
    
    # Lấy danh sách file .tex
    tex_files = [f for f in os.listdir(tex_folder) if f.endswith('.tex')]
    tex_files.sort()
    
    print(f"Found {len(tex_files)} LaTeX diagram files")
    
    # Biên dịch từng file .tex thành PDF đen trắng
    blackwhite_pdfs = []
    
    for tex_file in tex_files:
        tex_path = os.path.join(tex_folder, tex_file)
        pdf_name = tex_file.replace('.tex', '_bw.pdf')
        pdf_path = os.path.join(output_folder, pdf_name)
        
        try:
            # Đọc file LaTeX và thêm tùy chọn đen trắng
            with open(tex_path, 'r', encoding='utf-8') as f:
                tex_content = f.read()
            
            # Thêm dòng để tạo diagram đen trắng
            bw_tex_content = tex_content.replace(
                '\\begin{document}',
                '\\begin{document}\n\\pgfsetfillcolor{black}\n\\pgfsetstrokecolor{black}'
            )
            
            # Tạo file tạm
            temp_tex = os.path.join(output_folder, tex_file)
            with open(temp_tex, 'w', encoding='utf-8') as f:
                f.write(bw_tex_content)
            
            # Biên dịch LaTeX
            try:
                subprocess.run([
                    'pdflatex', 
                    '-interaction=nonstopmode',
                    '-output-directory', output_folder,
                    temp_tex
                ], check=True, capture_output=True)
                
                if os.path.exists(pdf_path):
                    blackwhite_pdfs.append(pdf_path)
                    print(f"Created: {pdf_name}")
                else:
                    # Tìm file PDF được tạo
                    for f in os.listdir(output_folder):
                        if f.endswith('.pdf') and f != pdf_name:
                            src = os.path.join(output_folder, f)
                            dst = os.path.join(output_folder, pdf_name)
                            if os.path.exists(src):
                                os.rename(src, dst)
                                blackwhite_pdfs.append(dst)
                                print(f"Created: {pdf_name}")
                            break
                            
            except subprocess.CalledProcessError as e:
                print(f"Failed to compile {tex_file}: {e}")
                # Sử dụng PDF gốc nếu không thể biên dịch
                original_pdf = tex_file.replace('.tex', '.pdf')
                original_pdf_path = os.path.join(tex_folder.replace('tex', 'pdf'), original_pdf)
                if os.path.exists(original_pdf_path):
                    shutil.copy(original_pdf_path, pdf_path)
                    blackwhite_pdfs.append(pdf_path)
                    print(f"Using original PDF: {pdf_name}")
                    
        except Exception as e:
            print(f"Error processing {tex_file}: {e}")
    
    print(f"Total black-white PDFs created: {len(blackwhite_pdfs)}")
    
    # Thêm vào file Word
    add_diagrams_to_word(blackwhite_pdfs, output_folder)
    
    # Dọn dẹp
    shutil.rmtree(output_folder)

def add_diagrams_to_word(pdf_files, pdf_folder):
    """Thêm các diagram PDF vào file Word du học sinh"""
    
    word_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
    
    doc = Document(word_doc)
    
    # Thêm section diagram
    diagram_heading = doc.add_heading('CÁC BIỂU ĐỒ THIẾT KẾ HỆ THỐNG', level=1)
    diagram_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Phân loại diagram
    activity_diagrams = [f for f in pdf_files if 'activity' in f.lower()]
    bfd_diagrams = [f for f in pdf_files if 'bfd' in f.lower()]
    dfd_diagrams = [f for f in pdf_files if 'dfd' in f.lower()]
    usecase_diagrams = [f for f in pdf_files if 'usecase' in f.lower()]
    class_diagrams = [f for f in pdf_files if 'class' in f.lower()]
    sequence_diagrams = [f for f in pdf_files if 'sequence' in f.lower()]
    state_diagrams = [f for f in pdf_files if 'state' in f.lower()]
    component_diagrams = [f for f in pdf_files if 'component' in f.lower()]
    erd_diagrams = [f for f in pdf_files if 'erd' in f.lower()]
    
    # Thêm từng loại diagram
    add_diagram_section(doc, 'ACTIVITY DIAGRAMS', activity_diagrams)
    add_diagram_section(doc, 'BUBBLE FLOW DIAGRAM', bfd_diagrams)
    add_diagram_section(doc, 'DATA FLOW DIAGRAMS', dfd_diagrams)
    add_diagram_section(doc, 'USECASE DIAGRAMS', usecase_diagrams)
    add_diagram_section(doc, 'CLASS DIAGRAMS', class_diagrams)
    add_diagram_section(doc, 'SEQUENCE DIAGRAMS', sequence_diagrams)
    add_diagram_section(doc, 'STATE DIAGRAMS', state_diagrams)
    add_diagram_section(doc, 'COMPONENT DIAGRAM', component_diagrams)
    add_diagram_section(doc, 'ENTITY RELATIONSHIP DIAGRAMS', erd_diagrams)
    
    # Lưu file
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_BLACKWHITE_DIAGRAMS.docx'
    doc.save(output_path)
    print(f"Document saved: {output_path}")

def add_diagram_section(doc, section_title, diagram_files):
    """Thêm một section diagram"""
    if not diagram_files:
        return
    
    # Thêm tiêu đề section
    heading = doc.add_heading(section_title, level=2)
    
    # Thêm từng diagram
    for pdf_file in diagram_files:
        pdf_name = os.path.basename(pdf_file).replace('_bw.pdf', '').replace('.pdf', '').replace('_', ' ').upper()
        diagram_heading = doc.add_heading(pdf_name, level=3)
        
        # Thêm placeholder (PDF không thể chèn trực tiếp vào Word)
        placeholder = doc.add_paragraph(f'[Black-white diagram: {os.path.basename(pdf_file)}]')
        placeholder.italic = True
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

if __name__ == '__main__':
    create_blackwhite_diagrams()
