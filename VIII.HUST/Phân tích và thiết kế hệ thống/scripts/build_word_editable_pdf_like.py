# -*- coding: utf-8 -*-
"""
Word EDITABLE + trình bày sát main.pdf (từ main.tex):
- Text từ LaTeX (đúng khoảng trắng tiếng Việt, sửa được)
- Mục lục + số trang từ PDF outline (y như PDF)
- Danh sách hình
- 27 sơ đồ PNG nét, bảng, danh sách
- TNR 12pt, lề 3/2/2/2 cm, header/footer như fancyhdr
Chỉ xuất: BaoCao/BaoCao_QuanLyKhoHang.docx
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

BASE = Path(r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao")
PDF_PATH = BASE / "main.pdf"
CHAPTERS_DIR = BASE / "chapters"
TEX_DIR = BASE / "diagrams" / "tex"
PNG_DIR = BASE / "diagrams" / "png"
OUTPUT = BASE / "BaoCao_QuanLyKhoHang.docx"

CHAPTERS = [
    "chapter1_khao_sat_hien_trang",
    "chapter2_mo_ta_nghiep_vu",
    "chapter3_phan_tich_chuc_nang",
    "chapter4_phan_tich_hanh_vi",
    "chapter5_phan_tich_tuong_tac",
    "chapter6_thiet_ke_lop_chi_tiet",
    "chapter7_thiet_ke_csdl_va_giao_dien",
]


# ── fonts / helpers ──────────────────────────────────────────

def font(run, size=12, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    font(run, 11)
    r = run._r
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r.append(b); r.append(i); r.append(e)


def extract_braced(s: str, start: int = 0):
    i = s.find("{", start)
    if i < 0:
        return "", start
    depth = 0
    for j in range(i, len(s)):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1 : j], j + 1
    return s[i + 1 :], len(s)


def clean_latex(text: str) -> str:
    if not text:
        return ""
    text = text.replace("~", " ").replace("\\,", " ")
    text = text.replace("\\%", "%").replace("\\&", "&").replace("\\$", "$")
    text = text.replace("\\_", "_").replace("\\#", "#")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace("\\ldots", "...").replace("\\dots", "...")
    text = text.replace("\\rightarrow", "→").replace("$\\rightarrow$", "→")
    text = re.sub(r"\$([^$]*)\$", r"\1", text)
    for _ in range(5):
        text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\underline\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\text\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\footnote\{([^{}]*)\}", r" (\1)", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def get_caption(name: str) -> str:
    p = TEX_DIR / f"{name}.tex"
    if not p.exists():
        return name
    t = p.read_text(encoding="utf-8", errors="ignore")
    m = re.search(r"\\caption\{", t)
    if not m:
        return name
    c, _ = extract_braced(t, m.start())
    return clean_latex(c) or name


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(12)
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    n.paragraph_format.line_spacing = 1.15
    n.paragraph_format.space_after = Pt(0)
    n.paragraph_format.space_before = Pt(0)

    for lvl, sz in ((1, 14), (2, 13), (3, 12)):
        try:
            st = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.space_before = Pt(14 if lvl == 1 else 10)
        st.paragraph_format.space_after = Pt(8 if lvl == 1 else 6)
        st.paragraph_format.line_spacing = 1.15

    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run("Phân tích và thiết kế hệ thống quản lý kho hàng")
    font(r, 10, italic=True)
    # header line
    pPr = hp._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)
    return doc


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font(r, {1: 14, 2: 13, 3: 12}[level], bold=True)
    # no first-line indent on headings
    p.paragraph_format.first_line_indent = Cm(0)
    return p


def add_body(doc, text, *, first_indent=True, justify=True):
    text = clean_latex(text)
    if not text:
        return None
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.0) if first_indent else Cm(0)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # bold segments: ** not used; handle "Mô tả:" style via simple split
    # Support leading Bold labels like "Mô tả: rest" when from \textbf{Mô tả}
    r = p.add_run(text)
    font(r, 12)
    return p


def add_rich_body(doc, raw: str):
    """Paragraph with \\textbf support kept as bold runs."""
    raw = raw.strip()
    if not raw:
        return
    # Convert remaining latex bold to markers
    parts = re.split(r"(\\textbf\{[^{}]*\})", raw)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for part in parts:
        if not part:
            continue
        m = re.match(r"\\textbf\{([^{}]*)\}", part)
        if m:
            t = clean_latex(m.group(1))
            if t:
                r = p.add_run(t)
                font(r, 12, bold=True)
        else:
            t = clean_latex(part)
            if t:
                r = p.add_run(t)
                font(r, 12)
    if not p.runs:
        doc._body._body.remove(p._p)


def add_list_item(doc, text, *, ordered=False, level=0):
    text = clean_latex(text)
    if not text:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(3)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.25 + 0.5 * level)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    bullet = "• " if not ordered else "– "
    # bold inside item
    parts = re.split(r"(\\textbf\{[^{}]*\})", text if "\\textbf" in text else text)
    # text already cleaned — just write with bullet
    r = p.add_run(bullet + text)
    font(r, 12)


def add_figure(doc, tex_name: str, fig_no: str):
    png = PNG_DIR / f"{tex_name}.png"
    cap = get_caption(tex_name)
    cap = re.sub(r"^Hình\s*[\d.]+\s*[:–\-]\s*", "", cap)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0)
    if png.exists():
        p.add_run().add_picture(str(png), width=Cm(15.2))
    else:
        r = p.add_run(f"[Thiếu hình: {tex_name}]")
        font(r, 11, italic=True)

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    c.paragraph_format.first_line_indent = Cm(0)
    r = c.add_run(f"Hình {fig_no}: {cap}")
    font(r, 11, italic=True)


def parse_table(block: str):
    block = re.sub(r"\\endfirsthead.*?\\endhead", "", block, flags=re.S)
    block = re.sub(r"\\endfoot.*?\\endlastfoot", "", block, flags=re.S)
    block = re.sub(r"\\hline", "", block)
    block = re.sub(r"\\cline\{[^}]*\}", "", block)
    block = re.sub(r"\\multicolumn\{\d+\}\{[^}]*\}\{([^{}]*)\}", r"\1", block)
    rows = []
    for raw in re.split(r"\\\\", block):
        raw = raw.strip()
        if not raw:
            continue
        cells = [clean_latex(c) for c in raw.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            font(r, 11, bold=(i == 0))
            p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()


def process_chapter(doc, path: Path, chap_num: int, fig_counter: dict):
    content = path.read_text(encoding="utf-8").replace("\r\n", "\n")
    content = re.sub(r"(?<!\\)%.*", "", content)
    pos, n = 0, len(content)
    sec_n = sub_n = 0

    while pos < n:
        if content[pos].isspace():
            pos += 1
            continue

        m = re.match(r"\\chapter\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            add_heading(doc, f"CHƯƠNG {chap_num}: {clean_latex(title)}", 1)
            sec_n = sub_n = 0
            pos = end
            continue

        m = re.match(r"\\section\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            sec_n += 1
            sub_n = 0
            add_heading(doc, f"{chap_num}.{sec_n} {clean_latex(title)}", 2)
            pos = end
            continue

        m = re.match(r"\\subsection\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            sub_n += 1
            add_heading(doc, f"{chap_num}.{sec_n}.{sub_n} {clean_latex(title)}", 3)
            pos = end
            continue

        m = re.match(r"\\input\{diagrams/tex/([^}]+)\}", content[pos:])
        if m:
            name = m.group(1).replace(".tex", "")
            key = f"c{chap_num}"
            fig_counter[key] = fig_counter.get(key, 0) + 1
            add_figure(doc, name, f"{chap_num}.{fig_counter[key]}")
            pos += m.end()
            continue

        m = re.match(r"\\begin\{([a-zA-Z*]+)\}(?:\[[^\]]*\])?(?:\{[^}]*\})?", content[pos:])
        if m:
            env = m.group(1).rstrip("*")
            start_env = pos + m.end()
            pat_b = re.compile(r"\\begin\{" + re.escape(env) + r"\*?\}")
            pat_e = re.compile(r"\\end\{" + re.escape(env) + r"\*?\}")
            depth, search, end_pos = 1, start_env, None
            while depth > 0 and search < n:
                mb, me = pat_b.search(content, search), pat_e.search(content, search)
                if me is None:
                    break
                if mb and mb.start() < me.start():
                    depth += 1
                    search = mb.end()
                else:
                    depth -= 1
                    if depth == 0:
                        end_pos = me.start()
                        pos = me.end()
                    else:
                        search = me.end()
            body = content[start_env:end_pos] if end_pos is not None else ""

            if env in ("itemize", "enumerate"):
                ordered = env == "enumerate"
                for item in re.split(r"\\item\b", body)[1:]:
                    item = item.strip()
                    nested = re.search(r"\\begin\{(enumerate|itemize)\}", item)
                    if nested:
                        before = item[: nested.start()].strip()
                        if before:
                            # keep bold markers until clean
                            add_list_item(doc, re.sub(r"\\textbf\{([^{}]*)\}", r"\1", before), ordered=ordered)
                        nenv = nested.group(1)
                        nm = re.search(
                            r"\\begin\{" + nenv + r"\}(.*?)\\end\{" + nenv + r"\}",
                            item, flags=re.S,
                        )
                        if nm:
                            for ni in re.split(r"\\item\b", nm.group(1))[1:]:
                                add_list_item(
                                    doc,
                                    re.sub(r"\\textbf\{([^{}]*)\}", r"\1", ni.strip()),
                                    ordered=(nenv == "enumerate"),
                                    level=1,
                                )
                    else:
                        add_list_item(doc, re.sub(r"\\textbf\{([^{}]*)\}", r"\1", item), ordered=ordered)
                continue

            if env in ("tabular", "longtable"):
                add_table(doc, parse_table(body))
                continue

            if env == "figure":
                im = re.search(r"\\input\{diagrams/tex/([^}]+)\}", body)
                if im:
                    name = im.group(1).replace(".tex", "")
                    key = f"c{chap_num}"
                    fig_counter[key] = fig_counter.get(key, 0) + 1
                    add_figure(doc, name, f"{chap_num}.{fig_counter[key]}")
                continue

            plain = clean_latex(body)
            if plain:
                add_body(doc, plain)
            continue

        if content[pos] == "\\":
            # skip lone commands
            m = re.match(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^{}]*\})*", content[pos:])
            if m:
                pos += m.end()
                continue
            pos += 1
            continue

        m = re.match(r"([^\\]+?)(?=\n\s*\n|\\[a-zA-Z]|\Z)", content[pos:], flags=re.S)
        if m:
            text = re.sub(r"\s*\n\s*", " ", m.group(1).strip())
            if text:
                # preserve bold in paragraph
                if "\\textbf" in text or "textbf" in content[pos : pos + len(m.group(0))]:
                    # re-read original slice for bold
                    raw_slice = content[pos : pos + m.end()]
                    raw_slice = re.sub(r"\s*\n\s*", " ", raw_slice.strip())
                    add_rich_body(doc, raw_slice)
                else:
                    add_body(doc, text)
            pos += m.end()
            continue
        pos += 1


def toc_line(doc, title: str, page, level: int):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm({1: 0, 2: 0.4, 3: 0.9}.get(level, 0))
    pf.tab_stops.add_tab_stop(Cm(15.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r = p.add_run(title)
    font(r, 12, bold=(level == 1))
    r2 = p.add_run(f"\t{page}")
    font(r2, 12)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    for text, sz in [
        ("ĐẠI HỌC BÁCH KHOA HÀ NỘI", 13),
        ("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG", 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        font(r, sz, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    for text, sz in [
        ("BÀI TẬP LỚN", 18),
        ("PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 16),
        ("QUẢN LÝ KHO HÀNG", 16),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        font(r, sz, bold=True)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhóm thực hiện")
    font(r, 12)
    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hà Nội, tháng 7 năm 2026")
    font(r, 12)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Mục lục", 1)
    if not PDF_PATH.exists():
        add_body(doc, "Không tìm thấy main.pdf — compile main.tex trước.", first_indent=False)
        doc.add_page_break()
        return

    toc = fitz.open(PDF_PATH).get_toc(simple=True)
    chap = sec = sub = 0
    for level, title, page in toc:
        title = title.strip()
        if level == 1:
            chap += 1
            sec = sub = 0
            toc_line(doc, f"{chap}  {title}", page, 1)
        elif level == 2:
            sec += 1
            sub = 0
            toc_line(doc, f"{chap}.{sec}  {title}", page, 2)
        else:
            sub += 1
            toc_line(doc, f"{chap}.{sec}.{sub}  {title}", page, 3)
    doc.add_page_break()


def add_lof(doc):
    add_heading(doc, "Danh sách hình vẽ", 1)
    for ci, ch in enumerate(CHAPTERS, 1):
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        seq = 0
        for m in re.finditer(r"\\input\{diagrams/tex/([^}]+)\}", text):
            name = m.group(1).replace(".tex", "")
            seq += 1
            cap = re.sub(r"^Hình\s*[\d.]+\s*[:–\-]\s*", "", get_caption(name))
            toc_line(doc, f"{ci}.{seq}  {cap}", "", 2)
    doc.add_page_break()


def ensure_pngs():
    PNG_DIR.mkdir(exist_ok=True)
    pdf_dir = BASE / "diagrams" / "pdf"
    n = 0
    for pdf in sorted(pdf_dir.glob("*.pdf")):
        if pdf.name.startswith("pdf_"):
            continue
        out = PNG_DIR / f"{pdf.stem}.png"
        if out.exists() and out.stat().st_mtime >= pdf.stat().st_mtime:
            n += 1
            continue
        doc = fitz.open(pdf)
        pix = doc[0].get_pixmap(matrix=fitz.Matrix(2.2, 2.2), alpha=False)
        pix.save(str(out))
        doc.close()
        n += 1
    return n


def verify(doc_path: Path) -> dict:
    d = Document(str(doc_path))
    imgs = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
    h1 = [p.text for p in d.paragraphs if p.style.name == "Heading 1"]
    h2 = sum(1 for p in d.paragraphs if p.style.name == "Heading 2")
    # sample editable text (not image)
    body = [p.text for p in d.paragraphs if p.text and p.style.name == "Normal"]
    info = {
        "mb": round(doc_path.stat().st_size / 1024 / 1024, 2),
        "paragraphs": len(d.paragraphs),
        "tables": len(d.tables),
        "images": imgs,
        "h1": h1,
        "h2": h2,
        "body_samples": body[3:6] if len(body) > 6 else body[:3],
        "has_toc": any("Mục lục" in t for t in h1),
        "has_lof": any("Danh sách hình" in t for t in h1),
        "chapters": sum(1 for t in h1 if t.startswith("CHƯƠNG")),
    }
    # checks
    assert imgs == 27, f"expected 27 figures, got {imgs}"
    assert info["chapters"] == 7, info["chapters"]
    assert info["has_toc"] and info["has_lof"]
    # Vietnamese spaces present in sample
    sample = " ".join(body[:20])
    assert "hệ thống" in sample or "Hệ thống" in sample, "missing proper Vietnamese spaces"
    return info


def main():
    print("=== Build EDITABLE Word (PDF-like) from main.tex ===")
    if not PDF_PATH.exists():
        print("ERROR: compile main.tex → main.pdf first")
        sys.exit(1)

    n_png = ensure_pngs()
    print(f"Diagrams PNG: {n_png}")

    doc = setup_doc()
    add_cover(doc)
    add_toc(doc)
    add_lof(doc)

    figs: dict = {}
    for i, ch in enumerate(CHAPTERS, 1):
        path = CHAPTERS_DIR / f"{ch}.tex"
        print(f"  CHƯƠNG {i}: {ch}")
        process_chapter(doc, path, i, figs)

    # only one output
    for old in BASE.glob("BaoCao_QuanLyKho*.docx"):
        if old.name != OUTPUT.name:
            try:
                old.unlink()
            except OSError:
                pass

    doc.save(str(OUTPUT))
    info = verify(OUTPUT)
    print()
    print("OUTPUT:", OUTPUT)
    print(f"Size: {info['mb']} MB | paras={info['paragraphs']} | tables={info['tables']} | images={info['images']}")
    print("H1:", info["h1"])
    print("H2 count:", info["h2"])
    print("Body sample:", info["body_samples"])
    print("CHECK OK — editable text + 27 figures + TOC/LOF like PDF")


if __name__ == "__main__":
    main()
