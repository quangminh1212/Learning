# -*- coding: utf-8 -*-
"""
Một file Word EDITABLE, trình bày theo main.pdf (LaTeX):
- Trang bìa, Mục lục (số trang từ PDF outline), Danh sách hình
- 7 chương + đủ 27 sơ đồ, bảng, danh sách
- Times New Roman 12pt, lề trái 3cm / còn lại 2cm, header/footer như PDF
Chỉ xuất: BaoCao/BaoCao_QuanLyKhoHang.docx
"""
from __future__ import annotations

import re
from pathlib import Path

import fitz
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips

BASE = Path(r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao")
PDF_PATH = BASE / "main.pdf"
CHAPTERS_DIR = BASE / "chapters"
TEX_DIR = BASE / "diagrams" / "tex"
PNG_DIR = BASE / "diagrams" / "png"
OUTPUT = BASE / "BaoCao_QuanLyKhoHang.docx"

CHAPTERS = [
    "chapter1_khao_sat_hien_trang",
    "chapter2_mo_ta_nghiep_vu",
    "chapter3_phan_tich_chuc_nang",
    "chapter4_phan_tich_hanh_vi",
    "chapter5_phan_tich_tuong_tac",
    "chapter6_thiet_ke_lop_chi_tiet",
    "chapter7_thiet_ke_csdl_va_giao_dien",
]

# Content width ≈ 21 - 3 - 2 = 16 cm
CONTENT_WIDTH = Cm(16.0)
FIG_WIDTH = Cm(15.5)


def set_run_font(run, size=12, bold=None, italic=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_field(paragraph):
    run = paragraph.add_run()
    set_run_font(run, size=11)
    r = run._r
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = " PAGE "
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    r.append(fc1)
    r.append(it)
    r.append(fc2)


def setup_doc() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    for level, size in ((1, 14), (2, 13), (3, 12)):
        try:
            st = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.paragraph_format.space_before = Pt(12 if level == 1 else 10)
        st.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        st.paragraph_format.line_spacing = 1.15

    # Header like fancyhdr
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hp.add_run("Phân tích và thiết kế hệ thống quản lý kho hàng")
    set_run_font(r, size=10, italic=True)

    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_field(fp)

    return doc


def extract_braced(s: str, start: int = 0) -> tuple[str, int]:
    i = s.find("{", start)
    if i < 0:
        return "", start
    depth = 0
    for j in range(i, len(s)):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1 : j], j + 1
    return s[i + 1 :], len(s)


def clean_latex(text: str) -> str:
    if not text:
        return ""
    text = text.replace("~", " ")
    text = text.replace("\\,", " ")
    text = text.replace("\\%", "%")
    text = text.replace("\\&", "&")
    text = text.replace("\\$", "$")
    text = text.replace("\\_", "_")
    text = text.replace("\\#", "#")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace("\\ldots", "...").replace("\\dots", "...")
    text = text.replace("\\rightarrow", "→").replace("$\\rightarrow$", "→")
    text = re.sub(r"\$([^$]*)\$", r"\1", text)
    # nested simple commands
    for _ in range(4):
        text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\underline\{([^{}]*)\}", r"\1", text)
        text = re.sub(r"\\text\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\footnote\{([^{}]*)\}", r" (\1)", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def get_caption(tex_name: str) -> str:
    path = TEX_DIR / f"{tex_name}.tex"
    if not path.exists():
        return tex_name
    t = path.read_text(encoding="utf-8", errors="ignore")
    m = re.search(r"\\caption\{", t)
    if not m:
        return tex_name
    content, _ = extract_braced(t, m.start())
    return clean_latex(content) or tex_name


def body_para(doc: Document, text: str, indent=True):
    text = clean_latex(text)
    if not text:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(1.0) if indent else Cm(0)
    r = p.add_run(text)
    set_run_font(r, size=12)
    return p


def heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_run_font(r, size={1: 14, 2: 13, 3: 12}.get(level, 12), bold=True)
    return p


def list_item(doc: Document, text: str, ordered=False, level=0):
    text = clean_latex(text)
    if not text:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(3)
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(1.0 + 0.5 * level)
    prefix = ("• " if not ordered else "– ")
    r = p.add_run(prefix + text)
    set_run_font(r, size=12)
    return p


def add_figure(doc: Document, tex_name: str, fig_no: str | None = None):
    png = PNG_DIR / f"{tex_name}.png"
    cap = get_caption(tex_name)
    if fig_no and not re.match(r"^Hình\s", cap, re.I):
        cap_text = f"Hình {fig_no}: {cap}"
    else:
        cap_text = cap if re.match(r"^Hình\s", cap, re.I) else f"Hình: {cap}"

    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(8)
    pf.space_after = Pt(4)
    pf.first_line_indent = Cm(0)
    if png.exists():
        run = p.add_run()
        run.add_picture(str(png), width=FIG_WIDTH)
    else:
        r = p.add_run(f"[Thiếu hình: {tex_name}]")
        set_run_font(r, size=11, italic=True)

    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    c.paragraph_format.first_line_indent = Cm(0)
    r = c.add_run(cap_text)
    set_run_font(r, size=11, italic=True)


def parse_table(block: str) -> list[list[str]]:
    block = re.sub(r"\\endfirsthead.*?\\endhead", "", block, flags=re.S)
    block = re.sub(r"\\endfoot.*?\\endlastfoot", "", block, flags=re.S)
    block = re.sub(r"\\hline", "", block)
    block = re.sub(r"\\cline\{[^}]*\}", "", block)
    block = re.sub(r"\\multicolumn\{\d+\}\{[^}]*\}\{([^{}]*)\}", r"\1", block)
    rows = []
    for raw in re.split(r"\\\\", block):
        raw = raw.strip()
        if not raw:
            continue
        cells = [clean_latex(c) for c in raw.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, txt in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(txt)
            set_run_font(r, size=11, bold=(i == 0))
    doc.add_paragraph()


def process_chapter(doc: Document, path: Path, chap_num: int, fig_counter: dict):
    content = path.read_text(encoding="utf-8")
    content = content.replace("\r\n", "\n")
    content = re.sub(r"(?<!\\)%.*", "", content)

    pos = 0
    n = len(content)
    sec_n = 0
    sub_n = 0

    while pos < n:
        if content[pos].isspace():
            pos += 1
            continue

        m = re.match(r"\\chapter\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            heading(doc, f"CHƯƠNG {chap_num}: {clean_latex(title)}", 1)
            sec_n = 0
            sub_n = 0
            pos = end
            continue

        m = re.match(r"\\section\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            sec_n += 1
            sub_n = 0
            heading(doc, f"{chap_num}.{sec_n} {clean_latex(title)}", 2)
            pos = end
            continue

        m = re.match(r"\\subsection\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos)
            sub_n += 1
            heading(doc, f"{chap_num}.{sec_n}.{sub_n} {clean_latex(title)}", 3)
            pos = end
            continue

        m = re.match(r"\\input\{diagrams/tex/([^}]+)\}", content[pos:])
        if m:
            name = m.group(1).replace(".tex", "")
            fig_counter["n"] = fig_counter.get("n", 0) + 1
            # figure number by chapter: chap.seq_in_chap
            key = f"c{chap_num}"
            fig_counter[key] = fig_counter.get(key, 0) + 1
            fig_no = f"{chap_num}.{fig_counter[key]}"
            add_figure(doc, name, fig_no=fig_no)
            pos += m.end()
            continue

        m = re.match(r"\\begin\{([a-zA-Z*]+)\}(?:\[[^\]]*\])?(?:\{[^}]*\})?", content[pos:])
        if m:
            env = m.group(1).rstrip("*")
            start_env = pos + m.end()
            pat_b = re.compile(r"\\begin\{" + re.escape(env) + r"\*?\}")
            pat_e = re.compile(r"\\end\{" + re.escape(env) + r"\*?\}")
            depth = 1
            search = start_env
            end_pos = None
            while depth > 0 and search < n:
                mb = pat_b.search(content, search)
                me = pat_e.search(content, search)
                if me is None:
                    break
                if mb and mb.start() < me.start():
                    depth += 1
                    search = mb.end()
                else:
                    depth -= 1
                    if depth == 0:
                        end_pos = me.start()
                        pos = me.end()
                    else:
                        search = me.end()
            body = content[start_env:end_pos] if end_pos is not None else ""

            if env in ("itemize", "enumerate"):
                ordered = env == "enumerate"
                for item in re.split(r"\\item\b", body)[1:]:
                    item = item.strip()
                    nested = re.search(r"\\begin\{(enumerate|itemize)\}", item)
                    if nested:
                        before = item[: nested.start()].strip()
                        if before:
                            list_item(doc, before, ordered=ordered)
                        nenv = nested.group(1)
                        nm = re.search(
                            r"\\begin\{" + nenv + r"\}(.*?)\\end\{" + nenv + r"\}",
                            item,
                            flags=re.S,
                        )
                        if nm:
                            for ni in re.split(r"\\item\b", nm.group(1))[1:]:
                                list_item(doc, ni.strip(), ordered=(nenv == "enumerate"), level=1)
                    else:
                        list_item(doc, item, ordered=ordered)
                continue

            if env in ("tabular", "longtable"):
                add_table(doc, parse_table(body))
                continue

            if env == "figure":
                im = re.search(r"\\input\{diagrams/tex/([^}]+)\}", body)
                if im:
                    name = im.group(1).replace(".tex", "")
                    key = f"c{chap_num}"
                    fig_counter[key] = fig_counter.get(key, 0) + 1
                    add_figure(doc, name, fig_no=f"{chap_num}.{fig_counter[key]}")
                continue

            plain = clean_latex(body)
            if plain:
                body_para(doc, plain)
            continue

        if content[pos] == "\\":
            # Unwrap common text formatting so content is not dropped
            mfmt = re.match(
                r"\\(textbf|textit|emph|texttt|underline|text)\*?\{([^{}]*)\}",
                content[pos:],
            )
            if mfmt:
                body_para(doc, mfmt.group(2), indent=False)
                pos += mfmt.end()
                continue
            m = re.match(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^{}]*\})*", content[pos:])
            if m:
                pos += m.end()
                continue
            pos += 1
            continue

        m = re.match(r"([^\\]+?)(?=\n\s*\n|\\[a-zA-Z]|\Z)", content[pos:], flags=re.S)
        if m:
            text = re.sub(r"\s*\n\s*", " ", m.group(1).strip())
            if text:
                body_para(doc, text)
            pos += m.end()
            continue
        pos += 1


def toc_line(doc: Document, title: str, page: int | str, level: int):
    """TOC line with dotted leader + page number (như LaTeX)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    pf.space_before = Pt(1)
    pf.line_spacing = 1.15
    pf.first_line_indent = Cm(0)
    # indent by level
    indent = {1: 0.0, 2: 0.5, 3: 1.0}.get(level, 0.0)
    pf.left_indent = Cm(indent)
    # tab stop near right margin of content (≈16cm)
    tab_pos = Cm(15.5)
    pf.tab_stops.add_tab_stop(tab_pos, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    r = p.add_run(title)
    set_run_font(r, size=12, bold=(level == 1))
    r2 = p.add_run(f"\t{page}")
    set_run_font(r2, size=12)


def add_cover(doc: Document):
    for _ in range(4):
        doc.add_paragraph()

    for text, size, bold in [
        ("ĐẠI HỌC BÁCH KHOA HÀ NỘI", 13, True),
        ("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG", 12, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)

    for _ in range(3):
        doc.add_paragraph()

    for text, size in [
        ("BÀI TẬP LỚN", 18),
        ("PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 16),
        ("QUẢN LÝ KHO HÀNG", 16),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        set_run_font(r, size=size, bold=True)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhóm thực hiện")
    set_run_font(r, size=12)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hà Nội, tháng 7 năm 2026")
    set_run_font(r, size=12)

    doc.add_page_break()


def add_toc_from_chapters(doc: Document):
    """Build TOC from current chapter .tex files (always in sync with content)."""
    heading(doc, "Mục lục", 1)

    for chap_num, ch in enumerate(CHAPTERS, start=1):
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            continue
        content = path.read_text(encoding="utf-8")
        content = re.sub(r"(?<!\\)%.*", "", content)

        m_ch = re.search(r"\\chapter\*?\s*\{", content)
        if m_ch:
            title, _ = extract_braced(content, m_ch.start())
            toc_line(doc, f"{chap_num}  {clean_latex(title)}", "", 1)

        sec = 0
        sub = 0
        pos = 0
        while pos < len(content):
            m_sec = re.search(r"\\section\*?\s*\{", content[pos:])
            m_sub = re.search(r"\\subsection\*?\s*\{", content[pos:])
            if m_sec is None and m_sub is None:
                break
            # pick whichever comes first
            if m_sec and (m_sub is None or m_sec.start() <= m_sub.start()):
                abs_pos = pos + m_sec.start()
                title, end = extract_braced(content, abs_pos)
                sec += 1
                sub = 0
                toc_line(doc, f"{chap_num}.{sec}  {clean_latex(title)}", "", 2)
                pos = end
            else:
                abs_pos = pos + m_sub.start()
                title, end = extract_braced(content, abs_pos)
                sub += 1
                toc_line(doc, f"{chap_num}.{sec}.{sub}  {clean_latex(title)}", "", 3)
                pos = end

    doc.add_page_break()


def add_toc_from_pdf(doc: Document):
    """Legacy: TOC from main.pdf outline. Prefer add_toc_from_chapters."""
    add_toc_from_chapters(doc)


def add_listoffigures(doc: Document):
    heading(doc, "Danh sách hình vẽ", 1)
    # Gather captions from tex files in chapter order
    fig_list = []
    for ci, ch in enumerate(CHAPTERS, start=1):
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        seq = 0
        for m in re.finditer(r"\\input\{diagrams/tex/([^}]+)\}", text):
            name = m.group(1).replace(".tex", "")
            seq += 1
            cap = get_caption(name)
            # strip leading "Hình ..." if already in caption
            cap = re.sub(r"^Hình\s*[\d.]+\s*[:–-]\s*", "", cap)
            fig_list.append((f"{ci}.{seq}", cap))

    for num, cap in fig_list:
        toc_line(doc, f"{num}  {cap}", "", 2)

    doc.add_page_break()


def main():
    print("Building single editable Word (match LaTeX PDF structure)...")
    # ensure pngs
    pngs = list(PNG_DIR.glob("*.png"))
    print(f"  diagrams PNG: {len(pngs)}")
    if len(pngs) < 20:
        print("  WARNING: few diagrams — run export first")

    doc = setup_doc()
    add_cover(doc)
    add_toc_from_pdf(doc)
    add_listoffigures(doc)

    fig_counter: dict = {}
    for i, ch in enumerate(CHAPTERS, start=1):
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            print(f"  missing {path}")
            continue
        print(f"  chapter {i}: {ch}")
        process_chapter(doc, path, i, fig_counter)

    # only one output
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    mb = OUTPUT.stat().st_size / (1024 * 1024)
    print(f"Saved ONLY: {OUTPUT}")
    print(f"Size: {mb:.2f} MB")


if __name__ == "__main__":
    main()
