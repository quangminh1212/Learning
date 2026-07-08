from docx import Document

def verify_converted_diagrams():
    doc_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_WITH_CONVERTED_DIAGRAMS.docx'
    
    doc = Document(doc_path)
    
    # Đếm hình ảnh
    image_count = 0
    for para in doc.paragraphs:
        if para._element.xpath('.//pic:pic'):
            image_count += 1
    
    # Lấy nội dung
    text_content = []
    for para in doc.paragraphs:
        text_content.append(para.text)
    
    full_text = '\n'.join(text_content)
    
    # Kiểm tra các section diagram
    diagram_sections = {
        'ACTIVITY DIAGRAMS': 'ACTIVITY DIAGRAMS' in full_text,
        'BUBBLE FLOW DIAGRAM': 'BUBBLE FLOW DIAGRAM' in full_text,
        'DATA FLOW DIAGRAMS': 'DATA FLOW DIAGRAMS' in full_text,
        'USECASE DIAGRAMS': 'USECASE DIAGRAMS' in full_text,
        'CLASS DIAGRAMS': 'CLASS DIAGRAMS' in full_text,
        'SEQUENCE DIAGRAMS': 'SEQUENCE DIAGRAMS' in full_text,
        'STATE DIAGRAMS': 'STATE DIAGRAMS' in full_text,
        'COMPONENT DIAGRAM': 'COMPONENT DIAGRAM' in full_text,
        'ENTITY RELATIONSHIP DIAGRAMS': 'ENTITY RELATIONSHIP DIAGRAMS' in full_text,
    }
    
    # Đếm số lượng diagram theo loại
    activity_count = full_text.count('ACTIVITY') if diagram_sections['ACTIVITY DIAGRAMS'] else 0
    usecase_count = full_text.count('USECASE') if diagram_sections['USECASE DIAGRAMS'] else 0
    class_count = full_text.count('CLASS') if diagram_sections['CLASS DIAGRAMS'] else 0
    sequence_count = full_text.count('SEQUENCE') if diagram_sections['SEQUENCE DIAGRAMS'] else 0
    
    with open(r'C:\Dev\Learning\converted_diagrams_verification.txt', 'w', encoding='utf-8') as f:
        f.write('CONVERTED DIAGRAMS VERIFICATION\n')
        f.write('================================\n\n')
        f.write(f'Total paragraphs: {len(text_content)}\n')
        f.write(f'Total characters: {len(full_text)}\n')
        f.write(f'Total images: {image_count}\n\n')
        
        f.write('DIAGRAM SECTIONS:\n')
        for section, present in diagram_sections.items():
            status = 'PRESENT' if present else 'MISSING'
            f.write(f'{section}: {status}\n')
        
        f.write('\nDIAGRAM COUNTS:\n')
        f.write(f'Activity diagrams: {activity_count}\n')
        f.write(f'Usecase diagrams: {usecase_count}\n')
        f.write(f'Class diagrams: {class_count}\n')
        f.write(f'Sequence diagrams: {sequence_count}\n')
        
        f.write('\nSAMPLE CONTENT (Last 20 lines):\n')
        for i, line in enumerate(text_content[-20:]):
            if line.strip():
                f.write(f'{len(text_content)-20+i+1}: {line[:80]}\n')
    
    print('Verification completed!')
    print('Results saved to converted_diagrams_verification.txt')
    print(f'Total images: {image_count}')
    print(f'Diagram sections present: {sum(diagram_sections.values())}/{len(diagram_sections)}')

if __name__ == '__main__':
    verify_converted_diagrams()
