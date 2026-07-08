# -*- coding: utf-8 -*-
"""
Xây dựng báo cáo Word đầy đủ từ LaTeX chapters + PNG diagrams.
Định dạng: Times New Roman 14pt, giãn dòng 1.5, lề trái 3cm, còn lại 2cm, A4.
"""
from __future__ import annotations

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, Twips, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

BASE = Path(r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao")
CHAPTERS_DIR = BASE / "chapters"
TEX_DIR = BASE / "diagrams" / "tex"
PNG_DIR = BASE / "diagrams" / "png"
OUTPUT = BASE / "BaoCao_QuanLyKhoHang.docx"

CHAPTERS = [
    "chapter1_khao_sat_hien_trang",
    "chapter2_mo_ta_nghiep_vu",
    "chapter3_phan_tich_chuc_nang",
    "chapter4_phan_tich_hanh_vi",
    "chapter5_phan_tich_tuong_tac",
    "chapter6_thiet_ke_lop_chi_tiet",
    "chapter7_thiet_ke_csdl_va_giao_dien",
]


def set_run_font(run, name="Times New Roman", size=14, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, first_line=True, space_after=6, space_before=0, align=None):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    if first_line:
        pf.first_line_indent = Cm(1.0)
    else:
        pf.first_line_indent = Cm(0)
    if align is not None:
        p.alignment = align


def add_page_number(paragraph):
    """Insert PAGE field."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    set_run_font(run, size=12)


def setup_document() -> Document:
    doc = Document()

    # A4 + margins (left 3cm, others 2cm)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    for level, size in [(1, 16), (2, 14), (3, 13)]:
        try:
            hs = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        hs.font.name = "Times New Roman"
        hs.font.size = Pt(size)
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(8)
        hs.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Header / Footer
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run("Phân tích và thiết kế hệ thống quản lý kho hàng")
    set_run_font(run, size=11, italic=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(fp)

    return doc


def clean_latex(text: str) -> str:
    if not text:
        return ""
    text = text.replace("~", " ")
    text = text.replace("\\,", " ")
    text = text.replace("\\%", "%")
    text = text.replace("\\&", "&")
    text = text.replace("\\$", "$")
    text = text.replace("\\_", "_")
    text = text.replace("\\#", "#")
    text = text.replace("\\{", "{")
    text = text.replace("\\}", "}")
    text = text.replace("``", "“")
    text = text.replace("''", "”")
    text = text.replace("\\ldots", "...")
    text = text.replace("\\dots", "...")
    text = text.replace("\\rightarrow", "→")
    text = text.replace("$\\rightarrow$", "→")
    text = text.replace("$\\Rightarrow$", "⇒")
    text = re.sub(r"\$([^$]+)\$", r"\1", text)
    # \textbf{...} \textit{...} etc keep content
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\underline\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\footnote\{([^{}]*)\}", r" (\1)", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extract_braced(s: str, start: int = 0) -> tuple[str, int]:
    """Extract content of first {...} starting at/after start. Returns (content, end_index)."""
    i = s.find("{", start)
    if i < 0:
        return "", start
    depth = 0
    for j in range(i, len(s)):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1 : j], j + 1
    return s[i + 1 :], len(s)


def get_caption_from_diagram(tex_name: str) -> str:
    path = TEX_DIR / f"{tex_name}.tex"
    if not path.exists():
        return tex_name
    text = path.read_text(encoding="utf-8", errors="ignore")
    m = re.search(r"\\caption\{", text)
    if not m:
        return tex_name
    content, _ = extract_braced(text, m.start())
    return clean_latex(content) or tex_name


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 16, 2: 14, 3: 13}.get(level, 14), bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_body_paragraph(doc: Document, text: str):
    text = clean_latex(text)
    if not text:
        return
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    # simple bold segments from leftover ** not used; just plain
    run = p.add_run(text)
    set_run_font(run, size=14)
    return p


def add_list_item(doc: Document, text: str, ordered: bool = False, level: int = 0):
    text = clean_latex(text)
    if not text:
        return
    style = "List Number" if ordered else "List Bullet"
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    p.clear()
    set_paragraph_format(p, first_line=False, space_after=3)
    p.paragraph_format.left_indent = Cm(1.0 + 0.5 * level)
    run = p.add_run(("• " if not ordered else "") + text)
    # for List Number style, numbering is automatic; avoid double bullet
    if ordered:
        p.clear()
        run = p.add_run(text)
    set_run_font(run, size=14)
    return p


def add_image(doc: Document, tex_name: str):
    png = PNG_DIR / f"{tex_name}.png"
    caption = get_caption_from_diagram(tex_name)
    if not png.exists():
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(f"[Thiếu hình: {tex_name}]")
        set_run_font(run, size=12, italic=True)
    else:
        # fit width ~ 15.5 cm (page content width ~16cm)
        p = doc.add_paragraph()
        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)
        run = p.add_run()
        run.add_picture(str(png), width=Cm(15.0))

    cap = doc.add_paragraph()
    set_paragraph_format(cap, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    run = cap.add_run(caption)
    set_run_font(run, size=12, italic=True)


def parse_table(block: str) -> list[list[str]]:
    """Parse tabular/longtable body into rows of cells."""
    # remove longtable headers noise
    block = re.sub(r"\\endfirsthead.*?\\endhead", "", block, flags=re.S)
    block = re.sub(r"\\endfoot.*?\\endlastfoot", "", block, flags=re.S)
    block = re.sub(r"\\hline", "", block)
    block = re.sub(r"\\cline\{[^}]*\}", "", block)
    block = re.sub(r"\\multicolumn\{[^}]*\}\{[^}]*\}\{([^}]*)\}", r"\1", block)
    rows = []
    for raw in re.split(r"\\\\", block):
        raw = raw.strip()
        if not raw or raw.startswith("\\"):
            continue
        cells = [clean_latex(c.strip()) for c in raw.split("&")]
        if any(cells):
            rows.append(cells)
    return rows


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < cols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    content_width = Cm(16.0)
    col_w = int(content_width / cols)
    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, size=12, bold=(row_idx == 0))
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
    doc.add_paragraph()


def process_chapter(doc: Document, path: Path):
    content = path.read_text(encoding="utf-8")
    # normalize newlines
    content = content.replace("\r\n", "\n")

    # remove comments
    content = re.sub(r"(?<!\\)%.*", "", content)

    pos = 0
    n = len(content)
    list_stack: list[str] = []  # 'itemize' | 'enumerate'

    while pos < n:
        # skip whitespace
        if content[pos].isspace():
            pos += 1
            continue

        # chapter
        m = re.match(r"\\chapter\*?\s*\{", content[pos:])
        if m:
            title, end = extract_braced(content, pos + m.start())
            add_heading(doc, "CHƯƠNG: " + clean_latex(title), 1)
            pos = end
            continue

        # section / subsection
        m = re.match(r"\\(sub)*section\*?\s*\{", content[pos:])
        if m:
            cmd = m.group(0)
            level = 2 if cmd.startswith("\\section") else 3
            title, end = extract_braced(content, pos)
            add_heading(doc, clean_latex(title), level)
            pos = end
            continue

        # input diagram
        m = re.match(r"\\input\{diagrams/tex/([^}]+)\}", content[pos:])
        if m:
            name = m.group(1).replace(".tex", "")
            add_image(doc, name)
            pos += m.end()
            continue

        # begin environments
        m = re.match(r"\\begin\{([a-zA-Z*]+)\}(?:\[[^\]]*\])?(?:\{[^}]*\})?", content[pos:])
        if m:
            env = m.group(1).rstrip("*")
            start_env = pos + m.end()
            # find matching end
            pattern_begin = re.compile(r"\\begin\{" + re.escape(env) + r"\*?\}")
            pattern_end = re.compile(r"\\end\{" + re.escape(env) + r"\*?\}")
            depth = 1
            search_pos = start_env
            end_pos = None
            while depth > 0 and search_pos < n:
                mb = pattern_begin.search(content, search_pos)
                me = pattern_end.search(content, search_pos)
                if me is None:
                    break
                if mb and mb.start() < me.start():
                    depth += 1
                    search_pos = mb.end()
                else:
                    depth -= 1
                    if depth == 0:
                        end_pos = me.start()
                        pos = me.end()
                    else:
                        search_pos = me.end()
            body = content[start_env:end_pos] if end_pos is not None else ""

            if env in ("itemize", "enumerate"):
                ordered = env == "enumerate"
                # split items
                items = re.split(r"\\item\b", body)
                for item in items[1:]:
                    # nested env - recursive simple: strip nested begin/end and keep text/items
                    item = item.strip()
                    # handle nested enumerate/itemize roughly
                    nested = re.search(r"\\begin\{(enumerate|itemize)\}", item)
                    if nested:
                        before = item[: nested.start()].strip()
                        if before:
                            add_list_item(doc, before, ordered=ordered)
                        # parse nested
                        nenv = nested.group(1)
                        nm = re.search(
                            r"\\begin\{" + nenv + r"\}(.*?)\\end\{" + nenv + r"\}",
                            item,
                            flags=re.S,
                        )
                        if nm:
                            for nitem in re.split(r"\\item\b", nm.group(1))[1:]:
                                add_list_item(doc, nitem.strip(), ordered=(nenv == "enumerate"), level=1)
                            after = item[nm.end() :].strip()
                            if after:
                                add_list_item(doc, after, ordered=ordered)
                    else:
                        add_list_item(doc, item, ordered=ordered)
                continue

            if env in ("tabular", "longtable"):
                rows = parse_table(body)
                add_table(doc, rows)
                continue

            if env == "figure":
                # look for input inside figure
                im = re.search(r"\\input\{diagrams/tex/([^}]+)\}", body)
                if im:
                    add_image(doc, im.group(1).replace(".tex", ""))
                else:
                    cm = re.search(r"\\caption\{", body)
                    if cm:
                        cap, _ = extract_braced(body, cm.start())
                        p = doc.add_paragraph()
                        set_paragraph_format(p, first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
                        run = p.add_run(clean_latex(cap))
                        set_run_font(run, size=12, italic=True)
                continue

            if env == "quote":
                add_body_paragraph(doc, body)
                continue

            # unknown env: treat body as text
            plain = clean_latex(body)
            if plain:
                add_body_paragraph(doc, plain)
            continue

        # item outside (shouldn't often happen)
        m = re.match(r"\\item\b\s*", content[pos:])
        if m:
            # read until newline or next command
            rest = content[pos + m.end() :]
            line = rest.split("\n", 1)[0]
            add_list_item(doc, line)
            pos += m.end() + len(line)
            continue

        # skip other commands
        if content[pos] == "\\":
            m = re.match(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^{}]*\})*", content[pos:])
            if m:
                pos += m.end()
                continue
            pos += 1
            continue

        # plain text paragraph until blank line or command
        m = re.match(r"([^\\]+?)(?=\n\s*\n|\\[a-zA-Z]|\Z)", content[pos:], flags=re.S)
        if m:
            text = m.group(1).strip()
            if text:
                # join lines
                text = re.sub(r"\s*\n\s*", " ", text)
                add_body_paragraph(doc, text)
            pos += m.end()
            continue

        pos += 1


def add_cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ĐẠI HỌC BÁCH KHOA HÀ NỘI")
    set_run_font(r, size=14, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TRƯỜNG CÔNG NGHỆ THÔNG TIN VÀ TRUYỀN THÔNG")
    set_run_font(r, size=13, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BÀI TẬP LỚN")
    set_run_font(r, size=18, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
    set_run_font(r, size=16, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QUẢN LÝ KHO HÀNG")
    set_run_font(r, size=16, bold=True)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Nhóm thực hiện")
    set_run_font(r, size=14)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Hà Nội, tháng 7 năm 2026")
    set_run_font(r, size=14)

    doc.add_page_break()

    # TOC placeholder
    add_heading(doc, "MỤC LỤC", 1)
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=False)
    r = p.add_run(
        "(Trong Microsoft Word: References → Table of Contents → Automatic Table "
        "để cập nhật mục lục theo các tiêu đề chương/mục.)"
    )
    set_run_font(r, size=12, italic=True)

    # simple TOC listing from chapters
    for ch in CHAPTERS:
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        for m in re.finditer(r"\\(chapter|section|subsection)\*?\{", text):
            title, _ = extract_braced(text, m.start())
            title = clean_latex(title)
            if not title:
                continue
            level = m.group(1)
            p = doc.add_paragraph()
            set_paragraph_format(p, first_line=False, space_after=2)
            indent = {"chapter": 0, "section": 0.5, "subsection": 1.0}[level]
            p.paragraph_format.left_indent = Cm(indent)
            prefix = "" if level != "chapter" else ""
            r = p.add_run(prefix + title)
            set_run_font(r, size=13 if level == "chapter" else 12, bold=(level == "chapter"))

    doc.add_page_break()


def main():
    print("Building full Word report...")
    print(f"PNG available: {len(list(PNG_DIR.glob('*.png')))}")
    doc = setup_document()
    add_cover(doc)

    for ch in CHAPTERS:
        path = CHAPTERS_DIR / f"{ch}.tex"
        if not path.exists():
            print(f"  MISSING {path}")
            continue
        print(f"  + {ch}")
        process_chapter(doc, path)

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
