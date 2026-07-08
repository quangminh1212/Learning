from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import traceback
import sys
import io

# Set encoding to UTF-8 for output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

try:
    input_file = r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\BaoCao_QuanLyKho_Full.docx"
    output_file = r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\BaoCao_QuanLyKho_Complete.docx"

    print(f"Processing {input_file}...")

    # Đọc tài liệu hiện tại
    doc = Document(input_file)

    # Tạo tài liệu mới
    new_doc = Document()

    # Copy styles
    new_doc.styles['Normal'].font.name = 'Times New Roman'
    new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Thêm trang bìa
    title = new_doc.add_heading('BÀI TẬP LỚN', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = new_doc.add_heading('PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    topic = new_doc.add_heading('QUẢN LÝ KHO HÀNG', level=2)
    topic.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = new_doc.add_paragraph('Nhóm thực hiện')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date = new_doc.add_paragraph('Hà Nội, tháng 7 năm 2026')
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_doc.add_page_break()

    # Thêm mục lục thủ công
    toc_heading = new_doc.add_heading('MỤC LỤC', level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mục lục cơ bản
    toc_items = [
        ("CHƯƠNG 1: KHẢO SÁT HIỆN TRẠNG", 1),
        ("1.1 Giới thiệu chung", 2),
        ("1.2 Đối tượng nghiên cứu", 2),
        ("1.3 Phương pháp nghiên cứu", 2),
        ("1.4 Kết luận chương", 2),
        ("CHƯƠNG 2: MÔ TẢ NGHIỆP VỤ", 3),
        ("2.1 Tổng quan các nghiệp vụ", 3),
        ("2.2 Nghiệp vụ quản lý người dùng", 3),
        ("2.3 Nghiệp vụ quản lý đối tác", 4),
        ("2.4 Nghiệp vụ quản lý hàng hóa và kho", 4),
        ("2.5 Nghiệp vụ quản lý nhập xuất", 5),
        ("2.6 Nghiệp vụ quản lý đơn hàng và tài chính", 6),
        ("2.7 Nghiệp vụ báo cáo", 6),
        ("CHƯƠNG 3: PHÂN TÍCH CHỨC NĂNG", 7),
        ("3.1 Phân tích tác nhân", 7),
        ("3.2 Sơ đồ phân cấp chức năng (BFD)", 7),
        ("3.3 Sơ đồ luồng dữ liệu (DFD)", 8),
        ("3.4 Sơ đồ use case", 9),
        ("3.5 Đặc tả use case", 10),
        ("CHƯƠNG 4: PHÂN TÍCH HÀNH VI", 13),
        ("4.1 Sơ đồ lớp tổng quát", 13),
        ("4.2 Sơ đồ lớp chi tiết theo nghiệp vụ", 14),
        ("CHƯƠNG 5: PHÂN TÍCH TƯƠNG TÁC", 17),
        ("5.1 Sơ đồ sequence", 17),
        ("5.2 Sơ đồ trạng thái", 20),
        ("CHƯƠNG 6: THIẾT KẾ LỚP CHI TIẾT", 21),
        ("6.1 Thiết kế lớp chi tiết", 21),
        ("6.2 Sơ đồ component", 24),
        ("CHƯƠNG 7: THIẾT KẾ CSDL VÀ GIAO DIỆN", 25),
        ("7.1 Thiết kế cơ sở dữ liệu", 25),
        ("7.2 Thiết kế giao diện", 27),
    ]

    for item, page in toc_items:
        para = new_doc.add_paragraph()
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5))
        
        run1 = para.add_run(item)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        
        # Add page number with tab
        run2 = para.add_run('\t')
        run3 = para.add_run(str(page))

    new_doc.add_page_break()

    # Thêm danh sách hình vẽ thủ công
    lof_heading = new_doc.add_heading('DANH SÁCH HÌNH VẼ', level=1)
    lof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Danh sách hình vẽ
    lof_items = [
        ("Hình 2.1 Sơ đồ hoạt động quy trình nhập kho", 11),
        ("Hình 2.2 Sơ đồ hoạt động quy trình xuất kho", 12),
        ("Hình 2.3 Sơ đồ hoạt động quy trình kiểm kê tồn kho", 13),
        ("Hình 3.1 Sơ đồ phân cấp chức năng (BFD)", 15),
        ("Hình 3.2 Sơ đồ luồng dữ liệu cấp 0 (Context Diagram)", 16),
        ("Hình 3.3 Sơ đồ luồng dữ liệu cấp 1", 17),
        ("Hình 3.4 Sơ đồ use case tổng quát", 18),
        ("Hình 3.5 Sơ đồ use case quản lý người dùng", 19),
        ("Hình 3.6 Sơ đồ use case quản lý đối tác", 19),
        ("Hình 3.7 Sơ đồ use case quản lý hàng hóa", 20),
        ("Hình 3.8 Sơ đồ use case quản lý tồn kho", 20),
        ("Hình 3.9 Sơ đồ use case báo cáo", 21),
        ("Hình 4.1 Sơ đồ lớp tổng quát", 28),
        ("Hình 4.2 Sơ đồ lớp nhập xuất kho", 29),
        ("Hình 4.3 Sơ đồ lớp kiểm kê và đơn hàng", 30),
        ("Hình 4.4 Sơ đồ lớp UC05", 31),
        ("Hình 4.5 Sơ đồ lớp UC09", 32),
        ("Hình 4.6 Sơ đồ lớp UC10", 32),
        ("Hình 5.1 Sơ đồ sequence UC05", 34),
        ("Hình 5.2 Sơ đồ sequence UC06", 35),
        ("Hình 5.3 Sơ đồ sequence UC09", 36),
        ("Hình 5.4 Sơ đồ sequence UC10", 37),
        ("Hình 5.5 Sơ đồ trạng thái phiếu nhập", 38),
        ("Hình 5.6 Sơ đồ trạng thái phiếu xuất", 38),
        ("Hình 6.1 Sơ đồ component", 43),
        ("Hình 7.1 Sơ đồ ERD master data", 45),
        ("Hình 7.2 Sơ đồ ERD transaction data", 45),
    ]

    for item, page in lof_items:
        para = new_doc.add_paragraph()
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.5))
        
        run1 = para.add_run(item)
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(12)
        
        # Add page number with tab
        run2 = para.add_run('\t')
        run3 = para.add_run(str(page))

    new_doc.add_page_break()

    # Copy nội dung từ tài liệu gốc (bỏ qua trang bìa)
    content_started = False
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Bỏ qua trang bìa
        if not content_started:
            if "CHƯƠNG 1" in text or "Khảo sát hiện trạng" in text:
                content_started = True
            else:
                continue
        
        if content_started:
            new_para = new_doc.add_paragraph(text)
            if para.style.name.startswith('Heading'):
                new_para.style = para.style
            elif para.runs and para.runs[0].bold:
                new_para.runs[0].bold = True

    # Copy tables
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                new_table.rows[i].cells[j].text = cell.text

    # Lưu tài liệu mới
    new_doc.save(output_file)
    print(f"Created complete document with TOC and LOF: {output_file}")

except Exception as e:
    print(f"Error during processing: {e}")
    traceback.print_exc()