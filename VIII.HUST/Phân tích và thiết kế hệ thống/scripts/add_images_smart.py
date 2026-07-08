from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import zipfile
import tempfile
import shutil

def extract_images_from_docx(doc_path, output_folder):
    """Trích xuất hình ảnh từ file docx"""
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    with zipfile.ZipFile(doc_path, 'r') as zip_ref:
        image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
        
        extracted_images = {}
        for img_file in image_files:
            img_data = zip_ref.read(img_file)
            img_name = os.path.basename(img_file)
            img_path = os.path.join(output_folder, img_name)
            
            with open(img_path, 'wb') as f:
                f.write(img_data)
            
            extracted_images[img_name] = img_path
    
    return extracted_images

def analyze_document_structure(doc_path):
    """Phân tích cấu trúc document chi tiết"""
    doc = Document(doc_path)
    
    structure = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        has_image = bool(para._element.xpath('.//pic:pic'))
        
        if text or has_image:
            structure.append({
                'index': i,
                'text': text,
                'has_image': has_image,
                'is_heading': para.style.name.startswith('Heading')
            })
    
    return structure

def add_images_to_correct_positions():
    """Thêm hình ảnh vào đúng vị trí"""
    
    # Đường dẫn file
    reference_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx'
    current_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_UPDATED.docx'
    
    # Thư mục tạm
    temp_folder = tempfile.mkdtemp()
    
    try:
        # Trích xuất hình ảnh từ tài liệu mẫu
        print("Extracting images from reference document...")
        ref_images_folder = os.path.join(temp_folder, 'reference_images')
        ref_images = extract_images_from_docx(reference_doc, ref_images_folder)
        print(f"Extracted {len(ref_images)} images")
        
        # Phân tích cấu trúc tài liệu mẫu
        print("Analyzing reference document structure...")
        ref_structure = analyze_document_structure(reference_doc)
        
        # Phân tích cấu trúc tài liệu hiện tại
        print("Analyzing current document structure...")
        curr_structure = analyze_document_structure(current_doc)
        
        # Mở tài liệu hiện tại
        doc = Document(current_doc)
        
        # Đọc tài liệu mẫu để lấy hình ảnh theo thứ tự
        ref_doc = Document(reference_doc)
        
        # Tìm các vị trí trong tài liệu mẫu có hình ảnh
        image_positions = []
        for i, para in enumerate(ref_doc.paragraphs):
            has_image = bool(para._element.xpath('.//pic:pic'))
            if has_image:
                # Lấy text trước hình ảnh
                text_before = ""
                for j in range(max(0, i-5), i):
                    if ref_doc.paragraphs[j].text.strip():
                        text_before = ref_doc.paragraphs[j].text.strip()
                        break
                
                image_positions.append({
                    'index': i,
                    'text_before': text_before
                })
        
        print(f"Found {len(image_positions)} image positions in reference document")
        
        # Thêm hình ảnh vào tài liệu hiện tại dựa trên context
        # Cách tiếp cận: tìm text tương tự trong tài liệu hiện tại và thêm hình ảnh sau đó
        
        added_count = 0
        image_list = list(ref_images.values())
        
        # Thêm hình ảnh vào các vị trí tương ứng
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Kiểm tra xem đây có phải là vị trí cần thêm hình ảnh không
            for img_pos in image_positions[:len(image_list)]:
                if text and img_pos['text_before'] and text == img_pos['text_before']:
                    # Thêm hình ảnh sau paragraph này
                    try:
                        # Tạo paragraph mới cho hình ảnh
                        img_para = doc.paragraphs[i].insert_paragraph_before("")
                        run = img_para.add_run()
                        
                        # Lấy hình ảnh tiếp theo
                        img_path = image_list[added_count % len(image_list)]
                        run.add_picture(img_path, width=Inches(6.0))
                        
                        # Căn giữa hình ảnh
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        added_count += 1
                        # print(f"Added image after: {text[:50]}...")
                        
                        if added_count >= len(image_list):
                            break
                    except Exception as e:
                        pass  # Error adding image
            
            if added_count >= len(image_list):
                break
        
        # Nếu chưa thêm đủ hình ảnh, thêm vào cuối document
        if added_count < len(image_list):
            # print(f"Adding remaining {len(image_list) - added_count} images to end of document...")
            pass
            for i in range(added_count, len(image_list)):
                try:
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(image_list[i], width=Inches(6.0))
                    added_count += 1
                    # print(f"Added image {added_count} to end")
                except Exception as e:
                    pass  # Error adding image to end
        
        # Lưu document
        output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_COMPLETE.docx'
        doc.save(output_path)
        print("Document saved successfully")
        print(f"Total images added: {added_count}")
        
        return output_path
        
    finally:
        # Dọn dẹp thư mục tạm
        shutil.rmtree(temp_folder)

if __name__ == '__main__':
    add_images_to_correct_positions()
