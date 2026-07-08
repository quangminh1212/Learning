from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_pdf_diagrams_to_word():
    """Thêm các diagram PDF vào file Word"""
    
    # File Word hiện tại
    word_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
    
    # Thư mục diagram PDF
    pdf_folder = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\diagrams\pdf'
    
    # Đọc file Word
    doc = Document(word_doc)
    
    # Lấy danh sách file PDF
    pdf_files = [f for f in os.listdir(pdf_folder) if f.endswith('.pdf')]
    pdf_files.sort()
    
    print(f"Found {len(pdf_files)} PDF diagrams")
    
    # Thêm tiêu đề section diagram
    diagram_heading = doc.add_heading('CÁC BIỂU ĐỒ THIẾT KẾ HỆ THỐNG', level=1)
    diagram_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Phân loại diagram
    activity_diagrams = [f for f in pdf_files if 'activity' in f]
    bfd_diagrams = [f for f in pdf_files if 'bfd' in f]
    dfd_diagrams = [f for f in pdf_files if 'dfd' in f]
    usecase_diagrams = [f for f in pdf_files if 'usecase' in f]
    class_diagrams = [f for f in pdf_files if 'class' in f]
    sequence_diagrams = [f for f in pdf_files if 'sequence' in f]
    state_diagrams = [f for f in pdf_files if 'state' in f]
    component_diagrams = [f for f in pdf_files if 'component' in f]
    erd_diagrams = [f for f in pdf_files if 'erd' in f]
    
    # Thêm từng loại diagram
    add_diagram_section(doc, 'ACTIVITY DIAGRAMS', activity_diagrams, pdf_folder)
    add_diagram_section(doc, 'BUBBLE FLOW DIAGRAM', bfd_diagrams, pdf_folder)
    add_diagram_section(doc, 'DATA FLOW DIAGRAMS', dfd_diagrams, pdf_folder)
    add_diagram_section(doc, 'USECASE DIAGRAMS', usecase_diagrams, pdf_folder)
    add_diagram_section(doc, 'CLASS DIAGRAMS', class_diagrams, pdf_folder)
    add_diagram_section(doc, 'SEQUENCE DIAGRAMS', sequence_diagrams, pdf_folder)
    add_diagram_section(doc, 'STATE DIAGRAMS', state_diagrams, pdf_folder)
    add_diagram_section(doc, 'COMPONENT DIAGRAM', component_diagrams, pdf_folder)
    add_diagram_section(doc, 'ENTITY RELATIONSHIP DIAGRAMS', erd_diagrams, pdf_folder)
    
    # Lưu file
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_FULL_DIAGRAMS.docx'
    doc.save(output_path)
    print(f"Document saved: {output_path}")

def add_diagram_section(doc, section_title, diagram_files, pdf_folder):
    """Thêm một section diagram"""
    if not diagram_files:
        return
    
    # Thêm tiêu đề section
    heading = doc.add_heading(section_title, level=2)
    
    # Thêm từng diagram
    for pdf_file in diagram_files:
        pdf_path = os.path.join(pdf_folder, pdf_file)
        
        # Thêm tiêu đề diagram
        diagram_name = pdf_file.replace('.pdf', '').replace('_', ' ').upper()
        diagram_heading = doc.add_heading(diagram_name, level=3)
        
        # Thêm placeholder cho PDF (Word không hỗ trợ trực tiếp PDF)
        # Trong thực tế, cần chuyển PDF thành hình ảnh trước
        placeholder = doc.add_paragraph(f'[Diagram: {pdf_file}]')
        placeholder.italic = True
        placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm dòng trống
        doc.add_paragraph()

if __name__ == '__main__':
    add_pdf_diagrams_to_word()
