from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import subprocess
import tempfile

def convert_pdf_to_images():
    """Chuyển PDF sang hình ảnh và thêm vào Word"""
    
    # Kiểm tra xem có ImageMagick hoặc công cụ khác không
    try:
        subprocess.run(['convert', '--version'], capture_output=True, check=True)
        has_imagemagick = True
    except:
        has_imagemagick = False
    
    # Kiểm tra pdf2image
    try:
        import pdf2image
        has_pdf2image = True
    except:
        has_pdf2image = False
    
    if not has_imagemagick and not has_pdf2image:
        # print("Cần cài đặt ImageMagick hoặc pdf2image để chuyển PDF sang hình ảnh")
        # print("Đang tạo file Word với placeholder...")
        create_word_with_placeholders()
        return
    
    # Nếu có công cụ chuyển đổi, thực hiện chuyển đổi
    if has_pdf2image:
        convert_with_pdf2image()
    else:
        convert_with_imagemagick()

def create_word_with_placeholders():
    """Tạo file Word với placeholder cho diagram"""
    
    word_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
    pdf_folder = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\diagrams\pdf'
    
    doc = Document(word_doc)
    pdf_files = [f for f in os.listdir(pdf_folder) if f.endswith('.pdf')]
    pdf_files.sort()
    
    # Thêm section diagram
    diagram_heading = doc.add_heading('CÁC BIỂU ĐỒ THIẾT KẾ HỆ THỐNG (THAM KHẢO)', level=1)
    diagram_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    note = doc.add_paragraph('Lưu ý: Các biểu đồ dưới đây từ dự án quản lý kho hàng để tham khảo cấu trúc')
    note.italic = True
    
    # Phân loại và thêm
    categories = {
        'ACTIVITY DIAGRAMS': [f for f in pdf_files if 'activity' in f],
        'BUBBLE FLOW DIAGRAM': [f for f in pdf_files if 'bfd' in f],
        'DATA FLOW DIAGRAMS': [f for f in pdf_files if 'dfd' in f],
        'USECASE DIAGRAMS': [f for f in pdf_files if 'usecase' in f],
        'CLASS DIAGRAMS': [f for f in pdf_files if 'class' in f],
        'SEQUENCE DIAGRAMS': [f for f in pdf_files if 'sequence' in f],
        'STATE DIAGRAMS': [f for f in pdf_files if 'state' in f],
        'COMPONENT DIAGRAM': [f for f in pdf_files if 'component' in f],
        'ENTITY RELATIONSHIP DIAGRAMS': [f for f in pdf_files if 'erd' in f],
    }
    
    for category, files in categories.items():
        if files:
            heading = doc.add_heading(category, level=2)
            for pdf_file in files:
                diagram_name = pdf_file.replace('.pdf', '').replace('_', ' ').upper()
                diagram_heading = doc.add_heading(diagram_name, level=3)
                
                placeholder = doc.add_paragraph(f'[File PDF: {pdf_file} - Cần chuyển thành hình ảnh]')
                placeholder.italic = True
                placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
    
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_WITH_DIAGRAM_PLACEHOLDERS.docx'
    doc.save(output_path)
    print("Document created with placeholders")

def convert_with_pdf2image():
    """Sử dụng pdf2image để chuyển PDF"""
    try:
        from pdf2image import convert_from_path
    except ImportError:
        # print("pdf2image not available")
        create_word_with_placeholders()
        return
    
    word_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_HOAN_CHINH.docx'
    pdf_folder = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao\diagrams\pdf'
    temp_folder = tempfile.mkdtemp()
    
    doc = Document(word_doc)
    pdf_files = [f for f in os.listdir(pdf_folder) if f.endswith('.pdf')]
    pdf_files.sort()
    
    # Thêm section diagram
    diagram_heading = doc.add_heading('CÁC BIỂU ĐỒ THIẾT KẾ HỆ THỐNG (THAM KHẢO)', level=1)
    diagram_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    categories = {
        'ACTIVITY DIAGRAMS': [f for f in pdf_files if 'activity' in f],
        'BUBBLE FLOW DIAGRAM': [f for f in pdf_files if 'bfd' in f],
        'DATA FLOW DIAGRAMS': [f for f in pdf_files if 'dfd' in f],
        'USECASE DIAGRAMS': [f for f in pdf_files if 'usecase' in f],
        'CLASS DIAGRAMS': [f for f in pdf_files if 'class' in f],
        'SEQUENCE DIAGRAMS': [f for f in pdf_files if 'sequence' in f],
        'STATE DIAGRAMS': [f for f in pdf_files if 'state' in f],
        'COMPONENT DIAGRAM': [f for f in pdf_files if 'component' in f],
        'ENTITY RELATIONSHIP DIAGRAMS': [f for f in pdf_files if 'erd' in f],
    }
    
    for category, files in categories.items():
        if files:
            heading = doc.add_heading(category, level=2)
            for pdf_file in files:
                pdf_path = os.path.join(pdf_folder, pdf_file)
                diagram_name = pdf_file.replace('.pdf', '').replace('_', ' ').upper()
                diagram_heading = doc.add_heading(diagram_name, level=3)
                
                try:
                    # Chuyển PDF sang hình ảnh
                    images = convert_from_path(pdf_path, dpi=150)
                    if images:
                        # Lấy trang đầu tiên
                        img_path = os.path.join(temp_folder, f'{pdf_file}.png')
                        images[0].save(img_path, 'PNG')
                        
                        # Thêm vào Word
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = para.add_run()
                        run.add_picture(img_path, width=Inches(6.0))
                except Exception as e:
                    placeholder = doc.add_paragraph(f'[Could not convert {pdf_file}: {e}]')
                    placeholder.italic = True
                
                doc.add_paragraph()
    
    output_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_WITH_CONVERTED_DIAGRAMS.docx'
    doc.save(output_path)
    print("Document created with converted diagrams")
    
    # Dọn dẹp
    import shutil
    shutil.rmtree(temp_folder)

def convert_with_imagemagick():
    """Sử dụng ImageMagick để chuyển PDF"""
    # print("ImageMagick conversion not implemented, using placeholders")
    create_word_with_placeholders()

if __name__ == '__main__':
    convert_pdf_to_images()
