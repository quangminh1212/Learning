# -*- coding: utf-8 -*-
"""
main.tex → main.pdf (đã compile) → Word
Ưu tiên Microsoft Word COM (mở PDF, lưu DOCX) để giữ layout.
Fallback: mỗi trang PDF = ảnh full A4 trong Word (visual y hệt PDF).
Chỉ xuất: BaoCao_QuanLyKhoHang.docx
"""
from __future__ import annotations

import io
import os
import shutil
import sys
import time
from pathlib import Path

import fitz
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

BASE = Path(r"C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\BaoCao")
PDF = BASE / "main.pdf"
OUT = BASE / "BaoCao_QuanLyKhoHang.docx"
TMP_WORD = BASE / "_tmp_word_convert.docx"

A4_W, A4_H = 21.0, 29.7
RENDER_SCALE = 2.5  # ~180 dpi, nét cho in


def convert_via_word_com(pdf: Path, out: Path) -> bool:
    """Dùng MS Word mở PDF và SaveAs DOCX."""
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        print("  [Word COM] pywin32 not available")
        return False

    pdf = pdf.resolve()
    out = out.resolve()
    tmp = TMP_WORD.resolve()
    if tmp.exists():
        tmp.unlink()

    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        print(f"  [Word COM] Opening PDF: {pdf}")
        # ConfirmConversions=False, ReadOnly=True, AddToRecentFiles=False
        doc = word.Documents.Open(
            str(pdf),
            False,  # ConfirmConversions
            True,   # ReadOnly
            False,  # AddToRecentFiles
        )
        time.sleep(1)
        # 16 = wdFormatXMLDocument (.docx)
        print(f"  [Word COM] SaveAs: {tmp}")
        doc.SaveAs2(str(tmp), FileFormat=16)
        doc.Close(False)
        doc = None
        word.Quit()
        word = None
        time.sleep(0.5)
        if not tmp.exists() or tmp.stat().st_size < 10_000:
            print("  [Word COM] output too small / missing")
            return False
        shutil.move(str(tmp), str(out))
        print(f"  [Word COM] OK → {out} ({out.stat().st_size/1024/1024:.2f} MB)")
        return True
    except Exception as e:
        print(f"  [Word COM] FAILED: {e}")
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        return False
    finally:
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def clear_para(p):
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Cm(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def convert_via_page_images(pdf: Path, out: Path) -> bool:
    """Mỗi trang PDF → 1 trang Word (ảnh full A4) — visual y hệt."""
    print(f"  [PageImage] Rendering {pdf} @ scale {RENDER_SCALE}")
    src = fitz.open(pdf)
    n = src.page_count
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(A4_W)
    sec.page_height = Cm(A4_H)
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.left_margin = sec.right_margin = sec.top_margin = sec.bottom_margin = Cm(0)
    sec.header_distance = Cm(0)
    sec.footer_distance = Cm(0)
    for h in (sec.header, sec.footer):
        for p in h.paragraphs:
            p.clear()

    mat = fitz.Matrix(RENDER_SCALE, RENDER_SCALE)
    for i in range(n):
        pix = src[i].get_pixmap(matrix=mat, alpha=False)
        bio = io.BytesIO(pix.tobytes("png"))
        if i == 0 and doc.paragraphs:
            p = doc.paragraphs[0]
            p.clear()
        else:
            p = doc.add_paragraph()
        clear_para(p)
        if i > 0:
            p.paragraph_format.page_break_before = True
        run = p.add_run()
        # Fit A4 exactly
        run.add_picture(bio, width=Cm(A4_W), height=Cm(A4_H))
        if (i + 1) % 10 == 0 or i + 1 == n:
            print(f"  [PageImage] {i+1}/{n}")
    src.close()
    doc.save(str(out))
    print(f"  [PageImage] OK → {out} ({out.stat().st_size/1024/1024:.2f} MB)")
    return True


def score_docx(path: Path, expected_pages: int) -> dict:
    """Đánh giá sơ bộ file Word."""
    info = {"path": str(path), "exists": path.exists(), "mb": 0, "images": 0, "paras": 0, "ok": False}
    if not path.exists():
        return info
    info["mb"] = round(path.stat().st_size / 1024 / 1024, 2)
    try:
        d = Document(str(path))
        info["paras"] = len(d.paragraphs)
        info["tables"] = len(d.tables)
        info["images"] = sum(1 for r in d.part.rels.values() if "image" in r.reltype)
        # good if many images (diagrams or pages) or substantial text
        info["ok"] = info["images"] >= min(20, expected_pages // 2) or info["paras"] > 100
    except Exception as e:
        info["error"] = str(e)
    return info


def visual_check(pdf: Path, docx: Path, sample_pages=(0, 1, 2, 5, 10)):
    """Render vài trang PDF ra PNG để kiểm tra nguồn; đếm ảnh trong Word."""
    out_dir = BASE / "_check_pages"
    out_dir.mkdir(exist_ok=True)
    src = fitz.open(pdf)
    for i in sample_pages:
        if i >= src.page_count:
            continue
        pix = src[i].get_pixmap(matrix=fitz.Matrix(1.2, 1.2), alpha=False)
        pix.save(str(out_dir / f"pdf_p{i+1:02d}.png"))
    src.close()
    print(f"  [Check] sample PDF pages → {out_dir}")
    sc = score_docx(docx, 53)
    print(f"  [Check] Word score: {sc}")
    return sc


def main():
    if not PDF.exists():
        print("ERROR: main.pdf missing — compile main.tex first")
        sys.exit(1)

    pdf = PDF.resolve()
    pages = fitz.open(pdf).page_count
    print(f"Source: {pdf} ({pages} pages)")

    # 1) Thử Word COM (editable hơn, layout Word import PDF)
    ok = convert_via_word_com(pdf, OUT)
    method = "WordCOM"

    # 2) Nếu Word COM kém (ít ảnh / file nhỏ) → page image (đẹp y hệt PDF)
    sc = score_docx(OUT, pages) if ok else {"ok": False, "images": 0, "mb": 0}
    print(f"  WordCOM score: {sc}")

    use_images = (not ok) or (sc.get("images", 0) < 15) or (sc.get("mb", 0) < 0.5)
    if use_images:
        print("  → Fallback / force PageImage for visual fidelity like PDF")
        convert_via_page_images(pdf, OUT)
        method = "PageImage"
        sc = score_docx(OUT, pages)

    sc = visual_check(pdf, OUT)
    print()
    print("=" * 50)
    print(f"METHOD : {method}")
    print(f"OUTPUT : {OUT}")
    print(f"SIZE   : {sc.get('mb')} MB | images={sc.get('images')} | paras={sc.get('paras')}")
    print(f"PDF    : {pages} pages")
    if method == "PageImage":
        print("NOTE   : Mỗi trang Word = 1 trang PDF (trình bày y hệt).")
        print("         Text là ảnh — in/nộp đẹp; sửa chữ nên sửa LaTeX rồi gen lại.")
    else:
        print("NOTE   : Word COM import — có thể chỉnh sửa; kiểm tra lại sơ đồ.")
    print("=" * 50)


if __name__ == "__main__":
    main()
