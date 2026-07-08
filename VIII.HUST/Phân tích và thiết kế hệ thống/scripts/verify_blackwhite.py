from docx import Document

def verify_blackwhite_document():
    doc_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_BAO_CAO_BLACKWHITE.docx'
    
    doc = Document(doc_path)
    
    # Đếm hình ảnh
    image_count = 0
    for para in doc.paragraphs:
        if para._element.xpath('.//pic:pic'):
            image_count += 1
    
    # Lấy nội dung
    text_content = []
    for para in doc.paragraphs:
        text_content.append(para.text)
    
    full_text = '\n'.join(text_content)
    
    # Kiểm tra các section
    sections = {
        'BLACK WHITE DIAGRAMS': 'MÀU ĐEN' in full_text or 'BLACK WHITE' in full_text,
        'ACTIVITY DIAGRAMS': 'ACTIVITY DIAGRAMS' in full_text,
        'USECASE DIAGRAMS': 'USECASE DIAGRAMS' in full_text,
        'CLASS DIAGRAMS': 'CLASS DIAGRAMS' in full_text,
        'SEQUENCE DIAGRAMS': 'SEQUENCE DIAGRAMS' in full_text,
        'SAMPLE DIAGRAMS': 'MẪU TỪ TÀI LIỆU CHUẨN' in full_text or 'SAMPLE' in full_text,
    }
    
    # Đếm số lượng diagram theo loại
    activity_count = full_text.count('activity') if sections['ACTIVITY DIAGRAMS'] else 0
    usecase_count = full_text.count('usecase') if sections['USECASE DIAGRAMS'] else 0
    class_count = full_text.count('class') if sections['CLASS DIAGRAMS'] else 0
    
    with open(r'C:\Dev\Learning\blackwhite_verification.txt', 'w', encoding='utf-8') as f:
        f.write('BLACK-WHITE DIAGRAMS VERIFICATION\n')
        f.write('==================================\n\n')
        f.write(f'Total paragraphs: {len(text_content)}\n')
        f.write(f'Total characters: {len(full_text)}\n')
        f.write(f'Total images: {image_count}\n\n')
        
        f.write('SECTIONS:\n')
        for section, present in sections.items():
            status = 'PRESENT' if present else 'MISSING'
            f.write(f'{section}: {status}\n')
        
        f.write('\nDIAGRAM TYPES:\n')
        f.write(f'Activity diagrams: {activity_count}\n')
        f.write(f'Usecase diagrams: {usecase_count}\n')
        f.write(f'Class diagrams: {class_count}\n')
        
        f.write('\nSAMPLE CONTENT (Last 15 lines):\n')
        for i, line in enumerate(text_content[-15:]):
            if line.strip():
                f.write(f'{len(text_content)-15+i+1}: {line[:80]}\n')
    
    print('Verification completed!')
    print('Results saved to blackwhite_verification.txt')
    print(f'Total images: {image_count}')
    print(f'Sections present: {sum(sections.values())}/{len(sections)}')

if __name__ == '__main__':
    verify_blackwhite_document()
