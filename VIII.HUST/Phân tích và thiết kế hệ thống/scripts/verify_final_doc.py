from docx import Document

def verify_final_document():
    doc_path = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh_COMPLETE.docx'
    
    doc = Document(doc_path)
    
    # Đếm hình ảnh
    image_count = 0
    for para in doc.paragraphs:
        if para._element.xpath('.//pic:pic'):
            image_count += 1
    
    # Đếm relationships
    rel_count = len(doc.part.rels) if doc.part.rels else 0
    
    # Đếm hình ảnh trong relationships
    image_rel_count = 0
    if doc.part.rels:
        for rel_id, rel in doc.part.rels.items():
            if 'image' in rel.target_ref:
                image_rel_count += 1
    
    # Kiểm tra các section quan trọng
    text_content = []
    for para in doc.paragraphs:
        text_content.append(para.text)
    
    full_text = '\n'.join(text_content)
    
    important_sections = {
        'PHAN CONG CONG VIEC': 'PHAN CONG CONG VIEC' in full_text.upper() or 'PHÂN CÔNG CÔNG VIỆC' in full_text,
        'KHAO SAT HIEN TRANG': 'KHAO SAT' in full_text.upper() or 'Khảo sát' in full_text,
        'THUC TRANG QUY TRINH': 'THUC TRANG' in full_text.upper() or 'Thực trạng' in full_text,
        'MUC TIEU HE THONG': 'MUC TIEU' in full_text.upper() or 'Mục tiêu' in full_text,
        'USECASE': 'USECASE' in full_text.upper() or 'usecase' in full_text.lower(),
        'CLASS DIAGRAM': 'CLASS' in full_text.upper() or 'Biểu đồ lớp' in full_text,
        'SEQUENCE DIAGRAM': 'SEQUENCE' in full_text.upper() or 'Biểu đồ trình tự' in full_text,
        'UC01': 'UC01' in full_text,
        'UC13': 'UC13' in full_text,
    }
    
    with open(r'C:\Dev\Learning\final_verification.txt', 'w', encoding='utf-8') as f:
        f.write('FINAL DOCUMENT VERIFICATION\n')
        f.write('============================\n\n')
        f.write(f'Total paragraphs: {len(text_content)}\n')
        f.write(f'Total characters: {len(full_text)}\n')
        f.write(f'Images in paragraphs: {image_count}\n')
        f.write(f'Total relationships: {rel_count}\n')
        f.write(f'Image relationships: {image_rel_count}\n\n')
        
        f.write('IMPORTANT SECTIONS:\n')
        for section, present in important_sections.items():
            status = 'PRESENT' if present else 'MISSING'
            f.write(f'{section}: {status}\n')
        
        f.write('\nALL SECTIONS PRESENT: ' + str(all(important_sections.values())) + '\n')
    
    print('Verification completed!')
    print('Results saved to final_verification.txt')
    print(f'Images in document: {image_count}')
    print(f'Image relationships: {image_rel_count}')
    
    all_present = all(important_sections.values())
    print(f'All important sections present: {all_present}')

if __name__ == '__main__':
    verify_final_document()
