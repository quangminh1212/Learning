from docx import Document
from docx.shared import Inches

def check_for_images(doc_path):
    doc = Document(doc_path)
    
    image_info = []
    
    # Kiểm tra từng đoạn văn
    for i, para in enumerate(doc.paragraphs):
        # Kiểm tra hình ảnh inline
        if para._element.xpath('.//pic:pic'):
            image_info.append(f"Paragraph {i}: Contains inline image")
    
    # Kiểm tra từng bảng
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    if para._element.xpath('.//pic:pic'):
                        image_info.append(f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {para_idx}: Contains image")
    
    # Kiểm tra relationships
    if doc.part.rels:
        image_info.append(f"Total relationships: {len(doc.part.rels)}")
        for rel_id, rel in doc.part.rels.items():
            if 'image' in rel.target_ref:
                image_info.append(f"Image relationship: {rel_id} -> {rel.target_ref}")
    
    return image_info

# Kiểm tra tài liệu mẫu
reference_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\TaiLieuMau\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh (2).docx'

print("Checking reference document for images...")
ref_images = check_for_images(reference_doc)

print(f"Found {len(ref_images)} image-related items in reference document:")
for info in ref_images:
    print(f"  {info}")

# Kiểm tra tài liệu hiện tại
current_doc = r'C:\Dev\Learning\VIII.HUST\Phân tích và thiết kế hệ thống\request\Bai Tap Lon\Nhóm 3\Nhóm 3_Phân tích&Thiết kế hệ thống quản lý du học sinh bậc đại học của trung tâm tư vấnn du học FlyHigh.docx'

print("\nChecking current document for images...")
curr_images = check_for_images(current_doc)

print(f"Found {len(curr_images)} image-related items in current document:")
for info in curr_images:
    print(f"  {info}")
